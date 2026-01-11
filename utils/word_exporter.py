import re
import os
from datetime import datetime
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.oxml import parse_xml
from docx.oxml.ns import nsmap


class ExamWordExporter:
    def __init__(self, template_type="exam_guangwai"):
        self.doc = Document()
        self.template_type = template_type

        # === 常量定义 ===
        self.FONT_CN_MAIN = '宋体'
        self.FONT_CN_BOLD = '黑体'
        self.FONT_EN = 'Times New Roman'

        # 字号映射 (Pt)
        self.SIZE_SMALL_2 = 18  # 小二
        self.SIZE_4 = 14  # 四号
        self.SIZE_SMALL_4 = 12  # 小四
        self.SIZE_5 = 10.5  # 五号
        self.SIZE_SMALL_5 = 9  # 小五

    # ================= 核心工具方法 =================

    def _set_font(self, run, size_pt, bold=False, font_name='宋体', align=None):
        """
        统一字体设置：
        1. 强制西文使用 Times New Roman
        2. 强制中文使用指定字体 (默认宋体)
        3. 统一字号和加粗
        """
        if not run: return

        # 基础属性
        run.font.name = self.FONT_EN
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        run.font.size = Pt(size_pt)
        run.font.bold = bold

        # 颜色统一为黑色 (防止有时出现自动色差异)
        run.font.color.rgb = RGBColor(0, 0, 0)

    def _clean_markdown(self, text):
        """
        深度清洗 Markdown 标记
        """
        if not text: return ""
        # 1. 去除标题标记 (#)
        text = re.sub(r'^#+\s*', '', text)
        # 2. 去除加粗/斜体 (**, __, *, _)
        text = re.sub(r'[\*\_]{1,2}', '', text)
        # 3. 去除列表标记 (-, +) 在行首的情况
        text = re.sub(r'^[\-\+]\s+', '', text)
        # 4. 去除代码块标记 (`)
        text = re.sub(r'`+', '', text)
        return text.strip()

    def _set_cell_borders(self, cell, top=False, bottom=False, left=False, right=False, sz=6, color="auto"):
        """
        使用底层 XML 设置单元格边框 (用于页脚特殊效果)
        """
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = tcPr.first_child_found_in("w:tcBorders")
        if tcBorders is None:
            tcBorders = OxmlElement('w:tcBorders')
            tcPr.append(tcBorders)

        for border_name, has_border in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
            if has_border:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), str(sz))  # 1/8 pt unit
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), color)
                tcBorders.append(border)

    def _add_page_number_field(self, paragraph):
        """插入系统生成的页码域: 第 PAGE 页 共 NUMPAGES 页"""
        self._set_font(paragraph.add_run("第 "), self.SIZE_SMALL_5)
        self._add_field(paragraph, "PAGE")
        self._set_font(paragraph.add_run(" 页 共 "), self.SIZE_SMALL_5)
        self._add_field(paragraph, "NUMPAGES")
        self._set_font(paragraph.add_run(" 页"), self.SIZE_SMALL_5)

    def _add_field(self, paragraph, field_code):
        run = paragraph.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = field_code
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._element.append(fldChar1)
        run._element.append(instrText)
        run._element.append(fldChar2)
        self._set_font(run, self.SIZE_SMALL_5)

    def _parse_big_questions(self, content):
        """解析大题，返回列表"""
        big_questions = []
        if not content: return big_questions

        # 匹配中文大题序号
        lines = content.split('\n')
        bq_pattern = re.compile(r'^\s*(#+\s*)?([一二三四五六七八九十]+[、\.])\s*(.*)')

        for line in lines:
            # 先清洗一下再匹配，防止 "**一、**" 这种情况干扰
            clean_line = self._clean_markdown(line)
            match = bq_pattern.match(clean_line)
            if match:
                marker = re.sub(r'[、\.]', '', match.group(2).strip())
                big_questions.append(marker)
        return big_questions

    # ================= 业务生成逻辑 =================

    def generate_guangwai_exam(self, content, metadata, output_path):
        # 1. 页面设置
        section = self.doc.sections[0]
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)  # 页脚底端距离
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.0)
        section.footer_distance = Cm(1.5)  # 设置页脚距离底部的距离

        # 2. 首页标题
        p_title = self.doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_title = p_title.add_run("广西外国语学院课程考核试卷")
        self._set_font(run_title, self.SIZE_SMALL_2, bold=True)

        # 3. 学年度信息
        now = datetime.now()
        year = now.year
        month = now.month
        if 2 <= month <= 7:
            y_start, y_end, semester = year - 1, year, "二"
        else:
            y_start, y_end, semester = year, year + 1, "一"

        p_year = self.doc.add_paragraph()
        p_year.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_year.paragraph_format.space_before = Pt(6)

        # 格式: 20 25 -- 20 26
        y1_p, y1_s = str(y_start)[:2], str(y_start)[2:]
        y2_p, y2_s = str(y_end)[:2], str(y_end)[2:]

        self._set_font(p_year.add_run(f"{y1_p} "), self.SIZE_4, bold=True)
        r = p_year.add_run(f"{y1_s} ")
        self._set_font(r, self.SIZE_4, bold=True);
        r.font.underline = True

        self._set_font(p_year.add_run(" — "), self.SIZE_4, bold=True)

        self._set_font(p_year.add_run(f"{y2_p} "), self.SIZE_4, bold=True)
        r = p_year.add_run(f"{y2_s} ")
        self._set_font(r, self.SIZE_4, bold=True);
        r.font.underline = True

        self._set_font(p_year.add_run(" 学年度第 "), self.SIZE_4, bold=True)
        r = p_year.add_run(f"  {semester}  ")
        self._set_font(r, self.SIZE_4, bold=True);
        r.font.underline = True
        self._set_font(p_year.add_run(" 学期"), self.SIZE_4, bold=True)

        # === 【插入密封线】 ===
        self._insert_seal_line_text_box(p_year)

        # 4. 考试类型
        p_type = self.doc.add_paragraph()
        p_type.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_type.paragraph_format.space_after = Pt(12)
        run_type = p_type.add_run("期末考试（ √ ）   补考（   ）    重新学习考试（   ）")
        self._set_font(run_type, self.SIZE_SMALL_4, bold=True)

        # 5. 表格 1 (课程信息) - 全加粗, 宋体五号
        table1 = self.doc.add_table(rows=5, cols=5)
        table1.style = 'Table Grid'
        table1.alignment = WD_TABLE_ALIGNMENT.CENTER
        table1.autofit = False

        # 列宽设置
        widths = [2.69, 4.94, 2.28, 2.28, 2.5]
        for row in table1.rows:
            row.height = Cm(0.9)
            for idx, cell in enumerate(row.cells):
                cell.width = Cm(widths[idx])
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # 填充内容辅助函数 (默认加粗, 宋体五号)
        def fill(r, c, text, merge_c=None):
            cell = table1.cell(r, c)
            if merge_c: cell.merge(table1.cell(r, merge_c))
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0] if p.runs else p.add_run(text)
            run.text = text  # 确保 text 被设置
            self._set_font(run, self.SIZE_5, bold=True)

        fill(0, 0, "课程名称")
        fill(0, 1, metadata.get('course_name', ''), merge_c=4)
        fill(1, 0, "学历层次")
        fill(1, 1, "本科（ √ ）/ 专科（ ）")
        fill(1, 2, "考核类型")
        fill(1, 3, "考查（ ）/ 考试（ √ ）", merge_c=4)
        fill(2, 0, "专业年级班级")
        fill(2, 1, metadata.get('class_name', ''))
        fill(2, 2, "考试时间")
        fill(2, 3, "（ 90 ）分钟", merge_c=4)
        fill(3, 0, "试卷类型")
        fill(3, 1, "开卷（ √ ）/ 闭卷（ ）")
        fill(3, 2, "命题教师")
        fill(3, 3, metadata.get('teacher_name', ''), merge_c=4)

        # 第五行特殊合并
        fill(4, 0, "系 (教研室)\n主任")
        fill(4, 1, metadata.get('head_name', ''))

        # 3-4列合并
        c_title = table1.cell(4, 2)
        c_title.merge(table1.cell(4, 3))
        p = c_title.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("二级学院（部）\n主管教学领导")
        self._set_font(r, self.SIZE_5, bold=True)

        fill(4, 4, "")  # 签字空

        self.doc.add_paragraph()  # 空行

        # 6. 表格 2 (得分汇总) - 逻辑修正
        big_qs = self._parse_big_questions(content)
        if not big_qs: big_qs = ["一", "二", "三", "四", "五"]

        num_q = len(big_qs)
        cols_count = 1 + num_q + 2  # 题号 + 题目数 + 总分 + 核分人
        table2 = self.doc.add_table(rows=3, cols=cols_count)
        table2.style = 'Table Grid'
        table2.alignment = WD_TABLE_ALIGNMENT.CENTER
        table2.autofit = False

        # 计算宽度
        w_fixed = 2.69 + 2.45 + 2.5  # 题号 + 总分 + 核分人
        w_remain = 14.69 - w_fixed
        w_q = w_remain / max(1, num_q)
        widths_2 = [2.69] + [w_q] * num_q + [2.45, 2.5]

        # 设置行高与列宽
        for r_idx, row in enumerate(table2.rows):
            row.height = Cm(0.8)
            for c_idx, cell in enumerate(row.cells):
                if c_idx < len(widths_2):
                    cell.width = Cm(widths_2[c_idx])
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # 填充辅助函数
        def fill2(r, c, text, is_bold=True):
            cell = table2.cell(r, c)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0] if p.runs else p.add_run(text)
            run.text = text
            self._set_font(run, self.SIZE_5, bold=is_bold)

        # === 第一行 (Row 0) ===
        fill2(0, 0, "题号")
        i = 0
        for i, q in enumerate(big_qs):
            fill2(0, i + 1, q)
        fill2(0, i+2, "总分")

        # 【修正点】明确设置第一行最后一列为“核分人”
        fill2(0, i+3, "核分人")

        # === 第二行 (Row 1) & 第三行 (Row 2) ===
        fill2(1, 0, "满分")
        fill2(2, 0, "实得分")

        # === 合并核分人下方的空白格 (供签字用) ===
        # 注意：只合并 Row 1 和 Row 2 的最后一列，Row 0 (表头) 保持独立
        c_sign_top = table2.cell(1, cols_count - 1)
        c_sign_bottom = table2.cell(2, cols_count - 1)
        c_sign_top.merge(c_sign_bottom)

        # self.doc.add_paragraph()

        # 7. 正文渲染
        self._render_exam_content(content)

        # 8. 页脚 (特殊构造)
        self._create_footer(section)

        self.doc.save(output_path)
        return output_path

    def _render_exam_content(self, content):
        """
        渲染试卷正文 (改进版：增加中文缩进，优化间距)
        """
        lines = content.split('\n')

        bq_pattern = re.compile(r'^\s*(#+\s*)?([一二三四五六七八九十]+[、\.])\s*(.*)')
        sq_pattern = re.compile(r'^\s*(\d+[\.\、])\s*(.*)')

        # 计算中文缩进 (五号字 10.5pt，2字符 ≈ 21pt)
        # 为了兼容性，使用 Pt 单位
        indent_2_chars = Pt(21)

        for line in lines:
            clean_text = self._clean_markdown(line)
            if not clean_text: continue

            bq_match = bq_pattern.match(clean_text)
            sq_match = sq_pattern.match(clean_text)

            if bq_match:
                # === 大题 ===
                self.doc.add_paragraph()  # 大题前空一行，区分度更高
                full_title = bq_match.group(2) + bq_match.group(3)
                self._add_question_block(full_title)

            elif sq_match:
                # === 小题 (1. xxxx) ===
                # 格式：首行缩进，行间距适中
                p = self.doc.add_paragraph()
                p.paragraph_format.line_spacing = 1.2
                p.paragraph_format.first_line_indent = indent_2_chars  # 首行缩进

                full_text = sq_match.group(1) + sq_match.group(2)
                run = p.add_run(full_text)
                self._set_font(run, self.SIZE_5, bold=False)

            else:
                # === 普通正文/选项 ===
                # 格式：首行缩进
                p = self.doc.add_paragraph()
                p.paragraph_format.line_spacing = 1.2
                p.paragraph_format.first_line_indent = indent_2_chars  # 首行缩进

                run = p.add_run(clean_text)
                self._set_font(run, self.SIZE_5, bold=False)

    def _remove_table_borders(self, table):
        """
        移除表格所有边框（用于外层布局容器）
        """
        for row in table.rows:
            for cell in row.cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcBorders = tcPr.first_child_found_in("w:tcBorders")
                if tcBorders is None:
                    tcBorders = OxmlElement('w:tcBorders')
                    tcPr.append(tcBorders)

                # 将上下左右及内部边框全部设为 nil
                for edge in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                    border = OxmlElement(f'w:{edge}')
                    border.set(qn('w:val'), 'nil')
                    tcBorders.append(border)

    def _insert_seal_line_text_box(self, paragraph):
        """
        使用 VML 插入左侧竖排密封线（精修版 V3）。
        改进点：
        1. 行间距改为固定值：第一行16磅，其他13磅。
        2. 重新计算虚线长度，防止溢出。
        """
        # 1. 注册命名空间
        from docx.oxml.ns import nsmap
        from docx.oxml import parse_xml
        if 'v' not in nsmap:
            nsmap['v'] = 'urn:schemas-microsoft-com:vml'
        if 'o' not in nsmap:
            nsmap['o'] = 'urn:schemas-microsoft-com:office:office'

        # 2. 准备文本内容
        u7 = "_" * 14
        u10 = "_" * 20
        sp = "   "

        # Line 1: 个人信息
        line1_text = f"姓名：{u7}{sp}学号：{u7}{sp}年级、专业、班级：{u10}{sp}座位号：{u7}"

        # Line 2: 密封线 (手动虚线 "- ")
        # 长度计算修正：A4侧边可用高度约 24-25cm。
        # "- " (2字符) 在 SimHei 小四中约占 3-4mm。
        # 之前计算可能过长，这里适当缩减单位数。

        dash_unit = "- "

        # 两端: 12个单位 (之前是18)
        side_unit_count = 8
        d_end = dash_unit * side_unit_count
        # 间隔: 5个单位 (之前是8)
        d_gap = dash_unit * 5

        chars = list("密封线内不要答题")
        joined_text = d_gap.join(chars)
        line2_text = f"{d_end}{joined_text}{d_end}"

        # Line 3: 纯虚线
        # 匹配 Line 2 的视觉长度
        # 8个汉字 ≈ 16个 dash_unit 长度
        # 总长 ≈ 12*2 + 5*7 + 16 = 24 + 35 + 16 = 75 units
        line3_text = dash_unit * (side_unit_count*2 + 5*7 + 8)

        # 3. 样式计算
        # 注意：高度设为 700pt (约24.7cm)，留出上下余量
        vml_style = (
            "position:absolute; "
            "mso-position-horizontal-relative:page; "
            "mso-position-vertical-relative:page; "
            "left:20pt; top:50pt; "
            "width:60pt; height:700pt; "
            "z-index:1"
        )

        # 4. 构建 VML XML (重点修改 w:spacing)
        # w:lineRule="exact" 表示固定值
        # w:line 单位是 1/20 pt。16pt = 320, 13pt = 260
        vml_xml = f"""
        <v:shape xmlns:v="urn:schemas-microsoft-com:vml" 
                 xmlns:o="urn:schemas-microsoft-com:office:office"
                 id="SealLineShape" 
                 style="{vml_style}"
                 filled="f" stroked="f" coordsize="21600,21600">

            <v:textbox style="layout-flow:vertical-ideographic; mso-layout-flow-alt:bottom-to-top" inset="0,0,0,0">
                <w:txbxContent xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">

                    <w:p>
                        <w:pPr>
                            <w:spacing w:line="320" w:lineRule="exact"/> 
                            <w:jc w:val="center"/> 
                        </w:pPr>
                        <w:r>
                            <w:rPr>
                                <w:rFonts w:ascii="SimHei" w:eastAsia="SimHei"/>
                                <w:sz w:val="24"/> 
                            </w:rPr>
                            <w:t xml:space="preserve">{line1_text}</w:t>
                        </w:r>
                    </w:p>

                    <w:p>
                        <w:pPr>
                            <w:spacing w:line="260" w:lineRule="exact" w:after="0"/> 
                            <w:jc w:val="center"/> 
                        </w:pPr>
                        <w:r>
                            <w:rPr>
                                <w:rFonts w:ascii="SimHei" w:eastAsia="SimHei"/>
                                <w:b/> <w:sz w:val="24"/>
                            </w:rPr>
                            <w:t xml:space="preserve">{line2_text}</w:t>
                        </w:r>
                    </w:p>

                    <w:p>
                        <w:pPr>
                            <w:spacing w:line="260" w:lineRule="exact"/>
                            <w:jc w:val="center"/> 
                        </w:pPr>
                        <w:r>
                            <w:rPr>
                                <w:rFonts w:ascii="SimHei" w:eastAsia="SimHei"/>
                                <w:b/> <w:sz w:val="24"/>
                            </w:rPr>
                            <w:t xml:space="preserve">{line3_text}</w:t>
                        </w:r>
                    </w:p>

                </w:txbxContent>
            </v:textbox>
        </v:shape>
        """

        run = paragraph.add_run()
        pict = parse_xml('<w:pict xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
        vml_element = parse_xml(vml_xml)
        pict.append(vml_element)
        run._element.append(pict)

    def _set_cell_specific_borders(self, cell, top='single', bottom='single', left='single', right='single'):
        """
        设置单元格的具体边框样式。
        参数值 'nil' 表示无边框，'single' 表示实线。
        """
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = tcPr.first_child_found_in("w:tcBorders")
        if tcBorders is None:
            tcBorders = OxmlElement('w:tcBorders')
            tcPr.append(tcBorders)

        for border_name, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
            # 查找或创建特定边框节点
            border = tcBorders.find(qn(f'w:{border_name}'))
            if border is None:
                border = OxmlElement(f'w:{border_name}')
                tcBorders.append(border)

            # 设置属性
            border.set(qn('w:val'), val)
            if val != 'nil':
                border.set(qn('w:sz'), '4')  # 边框粗细
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), 'auto')

    def _insert_seal_line(self, paragraph, img_path):
        """
        插入左侧密封线图片 (浮动定位)
        定位逻辑：
        - 水平：页面边距 (Page Margin) 左侧对齐。
               由于图片是在正文区浮动，我们需要计算绝对位置。
               目标：图片右边缘贴近左页边距(3.0cm)。
               公式：水平绝对位置 = 页边距(3.0cm) - 图片宽度。
        - 垂直：相对段落 (Academic Year) 下方对齐。
        """
        if not os.path.exists(img_path):
            print(f"[Warning] 密封线图片未找到: {img_path}")
            return

        run = paragraph.add_run()
        # 1. 插入图片对象 (此时是 inline)
        # 获取原始尺寸，假设图片高度适应页面，或者按比例缩放
        # 一般密封线图片高度较高，这里设定高度约 24cm (A4高29.7 - 上下边距)
        # 宽度自动缩放
        inline_shape = run.add_picture(img_path, height=Cm(24))

        # 获取图片 ID 和 文件名 (rId)
        inline = inline_shape._inline
        r_id = inline.graphic.graphicData.pic.blipFill.blip.embed

        # 获取图片实际尺寸 (EMU单位)
        extent = inline.extent
        cx = extent.cx
        cy = extent.cy

        # 计算水平位置：页边距 3cm - 图片宽度
        # 1 cm = 360000 EMU
        page_margin_left_emu = 3.0 * 360000
        pos_h_emu = int(page_margin_left_emu - cx)

        # 计算垂直位置：段落下方约 0.8cm (一行的高度)
        pos_v_emu = int(0.8 * 360000)

        # 2. 构造 Floating XML (Anchor)
        # 命名空间映射
        nsmap = {
            'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
            'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
            'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture'
        }

        # 构造 anchor 元素
        anchor = OxmlElement('wp:anchor')
        anchor.set('distT', "0")
        anchor.set('distB', "0")
        anchor.set('distL', "114300")
        anchor.set('distR', "114300")
        anchor.set('simplePos', "0")
        anchor.set('relativeHeight', "251658240")
        anchor.set('behindDoc', "1")  # 设为 1，位于文字下方，防止遮挡（虽在装订区，安全起见）
        anchor.set('locked', "0")
        anchor.set('layoutInCell', "1")
        anchor.set('allowOverlap', "1")

        # simplePos (占位)
        simplePos = OxmlElement('wp:simplePos')
        simplePos.set('x', "0")
        simplePos.set('y', "0")
        anchor.append(simplePos)

        # positionH (水平定位: relativeFrom="page")
        positionH = OxmlElement('wp:positionH')
        positionH.set('relativeFrom', "page")
        posOffsetH = OxmlElement('wp:posOffset')
        posOffsetH.text = str(pos_h_emu)
        positionH.append(posOffsetH)
        anchor.append(positionH)

        # positionV (垂直定位: relativeFrom="paragraph")
        # 这样可以跟随 "学年度" 那一行
        positionV = OxmlElement('wp:positionV')
        positionV.set('relativeFrom', "paragraph")
        posOffsetV = OxmlElement('wp:posOffset')
        posOffsetV.text = str(pos_v_emu)
        positionV.append(posOffsetV)
        anchor.append(positionV)

        # extent
        extent_el = OxmlElement('wp:extent')
        extent_el.set('cx', str(cx))
        extent_el.set('cy', str(cy))
        anchor.append(extent_el)

        # effectExtent
        effectExtent = OxmlElement('wp:effectExtent')
        effectExtent.set('l', "0")
        effectExtent.set('t', "0")
        effectExtent.set('r', "0")
        effectExtent.set('b', "0")
        anchor.append(effectExtent)

        # wrapNone (无环绕，避免影响正文布局)
        wrapNone = OxmlElement('wp:wrapNone')
        anchor.append(wrapNone)

        # docPr
        docPr = OxmlElement('wp:docPr')
        docPr.set('id', "666")  # 任意ID
        docPr.set('name', "Seal Line")
        anchor.append(docPr)

        # cNvGraphicFramePr
        cNvPr = OxmlElement('wp:cNvGraphicFramePr')
        aGraphicFrameLocks = OxmlElement('a:graphicFrameLocks')
        aGraphicFrameLocks.set(qn('a:noChangeAspect'), "1")
        cNvPr.append(aGraphicFrameLocks)
        anchor.append(cNvPr)

        # graphic (直接复用 inline 中的 graphic)
        anchor.append(inline.graphic)

        # 3. 替换 XML
        # inline.getparent() 是 <w:r> 的子元素 <w:drawing>
        # 我们要用 <wp:anchor> 替换 <wp:inline>
        drawing = inline.getparent()
        drawing.replace(inline, anchor)

    def _add_question_block(self, title_text):
        """
        插入大题结构 (无缝嵌套版)
        布局：
        1. 外层表格 (1行2列，无边框)：
           - 左列：垂直顶对齐 (TOP)，边距为0，嵌套得分表格。
           - 右列：垂直居中对齐 (CENTER)，放置标题。
        """
        # === 1. 外层容器表格 ===
        outer_tbl = self.doc.add_table(rows=1, cols=2)
        outer_tbl.autofit = False
        self._remove_table_borders(outer_tbl)  # 清除外层边框

        # 设置外层列宽
        w_container_left = Cm(3.69)
        w_container_right = Cm(11.0)

        cell_outer_left = outer_tbl.cell(0, 0)
        cell_outer_right = outer_tbl.cell(0, 1)

        cell_outer_left.width = w_container_left
        cell_outer_right.width = w_container_right

        # === 关键调整 1: 对齐方式 ===
        # 左侧：顶端对齐，确保小表格靠上
        cell_outer_left.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        # 右侧：居中对齐，确保标题文字垂直居中
        cell_outer_right.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # === 关键调整 2: 清除左侧单元格边距 ===
        # 让内部的得分表格能紧贴左上角，消除间隙
        self._set_cell_margins_zero(cell_outer_left)

        # === 2. 左侧：嵌套得分小表格 (2行2列) ===
        # 注意：add_table 默认会添加在单元格现有段落之后。
        # 由于我们清除了边距，它会紧贴顶部。
        score_tbl = cell_outer_left.add_table(rows=2, cols=2)
        score_tbl.style = 'Table Grid'
        score_tbl.autofit = False

        # 设置内层表格行高
        score_tbl.rows[0].height = Cm(0.81)
        score_tbl.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        score_tbl.rows[1].height = Cm(1.03)
        score_tbl.rows[1].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

        # 设置内层表格列宽 (总宽 3.69cm，填满父容器)
        w_inner_0 = Cm(1.44)
        w_inner_1 = Cm(2.25)

        for row in score_tbl.rows:
            row.cells[0].width = w_inner_0
            row.cells[1].width = w_inner_1
            row.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            row.cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # 填充得分表头
        def set_score_text(r, c, txt):
            cell = score_tbl.cell(r, c)
            # 清理单元格默认段落
            paragraph = cell.paragraphs[0]
            paragraph.clear()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 这里的 run 处理防止样式跑偏
            run = paragraph.add_run(txt)
            self._set_font(run, self.SIZE_SMALL_4, bold=True, font_name=self.FONT_CN_BOLD)

        set_score_text(0, 0, "得分")
        set_score_text(0, 1, "评卷人")

        # === 3. 右侧：大题标题 ===
        # 确保右侧单元格有段落
        p_title = cell_outer_right.paragraphs[0]
        p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # 设置左缩进：一个中文字符宽度 (保持设计感)
        p_title.paragraph_format.left_indent = Pt(12)

        run_title = p_title.add_run(title_text)
        self._set_font(run_title, self.SIZE_SMALL_4, bold=True)

    def _set_cell_margins_zero(self, cell):
        """
        [XML注入] 强制将单元格的内部边距 (Margins) 设置为 0。
        用于让嵌套表格完全贴合父单元格的边缘。
        """
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcMar = tcPr.first_child_found_in("w:tcMar")
        if tcMar is None:
            tcMar = OxmlElement('w:tcMar')
            tcPr.append(tcMar)

        # 将 上下左右 边距全部设为 0 (单位 dxa)
        for edge in ['top', 'left', 'bottom', 'right']:
            node = tcMar.find(qn(f'w:{edge}'))
            if node is None:
                node = OxmlElement(f'w:{edge}')
                tcMar.append(node)
            node.set(qn('w:w'), '0')
            node.set(qn('w:type'), 'dxa')

    def _create_footer(self, section):
        """
        页脚构造（修复 BlockItemContainer.add_table 缺少 width 参数的 bug）
        """
        footer = section.footer
        p = footer.paragraphs[0]
        p.clear()

        # 计算表格总宽度：页面宽度 - 左边距 - 右边距
        # 21 - 3 - 2 = 16 cm
        table_width = Cm(16.0)

        # 【修复点】这里必须显式传入 width 参数
        tbl = footer.add_table(rows=1, cols=3, width=table_width)
        tbl.autofit = False
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 列宽分配
        mid_w = Cm(3.0)
        side_w = Cm((16.0 - 3.0) / 2)  # 6.5 cm

        c_left = tbl.cell(0, 0)
        c_mid = tbl.cell(0, 1)
        c_right = tbl.cell(0, 2)

        c_left.width = side_w
        c_mid.width = mid_w
        c_right.width = side_w

        # 设置边框
        self._set_cell_borders(c_left, top=True)
        self._set_cell_borders(c_right, top=True)
        self._set_cell_borders(c_mid, top=False)

        # 填充内容
        p_left = c_left.paragraphs[0]
        p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_left.paragraph_format.space_before = Pt(4)
        run_l = p_left.add_run("广西外国语学院课程考核试卷")
        self._set_font(run_l, self.SIZE_SMALL_5)

        p_mid = c_mid.paragraphs[0]
        p_mid.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_mid.paragraph_format.space_before = Pt(4)
        self._add_page_number_field(p_mid)

        p_right = c_right.paragraphs[0]
        p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_right.paragraph_format.space_before = Pt(4)

        spacer = "\u00A0" * 13
        text_r = f"{spacer}考试过程中不得将试卷拆开"
        run_r = p_right.add_run(text_r)
        self._set_font(run_r, self.SIZE_SMALL_5)