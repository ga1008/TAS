import os
import re
from datetime import datetime
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.shared import Cm, Pt
from .word_exporter_base import BaseWordExporter


class GuangWaiStandardExporter(BaseWordExporter):
    def generate(self, content, metadata, output_path):
        # 1. 页面设置
        section = self.setup_page(1.5, 1.5, 1.5, 1.5, footer_distance=1.75)

        # 2. 标题
        p_title = self.doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.set_font(p_title.add_run("广西外国语学院课程考核评分细则"), self.SIZE_SMALL_2, bold=True)

        # 3. 学年度信息 (改进点：接收前端传参，精准控制下划线格式)
        self._create_semester_header(metadata)

        # 4. 考核说明
        note_text = metadata.get('assessment_note', '（非笔试考核）')
        p_note = self.doc.add_paragraph()
        p_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.set_font(p_note.add_run(note_text), self.SIZE_SMALL_4)

        # 5. 表格1 (基本信息)
        self._create_info_table(metadata)

        # 6. "评分细则" 标题
        self.doc.add_paragraph()  # 空行
        p_header = self.doc.add_paragraph()
        self.set_font(p_header.add_run("评分细则"), self.SIZE_5, bold=True)

        # 7. 评分细则正文 (改进点：深度解析 Markdown，智能排版)
        self._create_content_table(content)

        # 8. 底部提示信息
        self._create_footer_notes()

        # 9. 页脚页码
        self._create_footer(section)

        self.doc.save(output_path)
        return output_path

    def _create_semester_header(self, metadata):
        """
        生成：20 25 - 20 26 学年度第 一 学期
        要求：年份前两位不加下划线，后两位加空格和下划线
        """
        # 获取前端传来的值，如果为空则由后端兜底计算
        y_start = str(metadata.get('year_start', ''))
        y_end = str(metadata.get('year_end', ''))
        semester = str(metadata.get('semester', ''))

        if not (y_start and y_end and semester):
            y_start_calc, y_end_calc, semester_calc = self.get_semester_info()
            if not y_start: y_start = y_start_calc
            if not y_end: y_end = y_end_calc
            if not semester: semester = semester_calc

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)

        # 辅助渲染年份的函数：拆分 20 25
        def add_year_runs(year_str):
            if len(year_str) == 4:
                prefix = year_str[:2]  # 20
                suffix = year_str[2:]  # 25
            else:
                prefix = ""
                suffix = year_str

            # 前缀：正常
            self.set_font(p.add_run(prefix), self.SIZE_4, bold=True)
            # 后缀：加空格，加下划线
            r = p.add_run(f" {suffix} ")
            self.set_font(r, self.SIZE_4, bold=True, underline=True)

        self.set_font(p.add_run("（ "), self.SIZE_4, bold=True)

        add_year_runs(y_start)
        self.set_font(p.add_run(" － "), self.SIZE_4, bold=True)
        add_year_runs(y_end)

        self.set_font(p.add_run(" 学年度第 "), self.SIZE_4, bold=True)

        # 学期
        r_sem = p.add_run(f" {semester} ")
        self.set_font(r_sem, self.SIZE_4, bold=True, underline=True)

        self.set_font(p.add_run(" 学期）"), self.SIZE_4, bold=True)

    def _create_content_table(self, content):
        """
        创建正文表格，并进行 Markdown -> Word 的智能转换
        """
        table = self.doc.add_table(rows=1, cols=1)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        table.allow_autofit = False

        cell = table.cell(0, 0)
        cell.width = Cm(17.5)
        cell.paragraphs[0].text = ""  # 清空

        # 预处理：按行分割
        lines = content.split('\n')

        # 正则预编译
        # 1. 过滤垃圾分隔符 (如 *---, ---, ===)
        re_garbage = re.compile(r'^[\*\-\=]{3,}\s*$')
        # 2. 识别 Markdown 标题 (#)
        re_md_heading = re.compile(r'^(#+)\s*(.*)')
        # 3. 识别中文大纲 (一、xxx, 1. xxx)
        re_cn_heading_1 = re.compile(r'^\s*([一二三四五六七八九十]+[、\.])\s*(.*)')
        re_cn_heading_2 = re.compile(r'^\s*(\d+[\.\、])\s*(.*)')
        # 4. 识别列表符 (- xxx, * xxx)
        re_list_item = re.compile(r'^\s*[\-\*]\s+(.*)')
        # 5. 识别加粗 (**xxx**)
        re_bold = re.compile(r'(\*\*(.*?)\*\*)')

        for line in lines:
            line = line.strip()
            if not line: continue
            if re_garbage.match(line): continue  # 跳过分隔符

            # 初始化样式状态
            is_bold_line = False
            indent_level = 0  # 0=首行缩进, 1=悬挂缩进(标题)
            clean_text = line

            # --- 结构解析 ---

            # Case A: Markdown 标题 (#)
            match_md = re_md_heading.match(line)
            if match_md:
                level = len(match_md.group(1))
                clean_text = match_md.group(2)
                is_bold_line = True  # 标题整行加粗
                indent_level = 1  # 标题无需首行缩进，可能需要段前间距

            # Case B: 中文一级标题 (一、)
            elif re_cn_heading_1.match(line):
                is_bold_line = True
                indent_level = 1

            # Case C: 中文二级标题 (1.)
            elif re_cn_heading_2.match(line):
                # 视情况而定，通常小标题加粗
                is_bold_line = len(line) < 30  # 如果太长可能是正文列表，短的才是标题
                indent_level = 0  # 保持首行缩进

            # Case D: 列表项 (- )
            elif re_list_item.match(line):
                clean_text = re_list_item.match(line).group(1)
                indent_level = 0  # 列表项通常也需要缩进

            # --- 写入段落 ---
            p = cell.add_paragraph()

            # 设置段落格式
            if indent_level == 0:
                # 正文：首行缩进2字符 (五号字约10.5pt，2字符约21pt)
                p.paragraph_format.first_line_indent = Pt(21)
                p.paragraph_format.line_spacing = 1.25
            else:
                # 标题：左对齐，无首行缩进，稍微增加段间距
                p.paragraph_format.first_line_indent = 0
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(4)
                # 确保行距适应大号字
                p.paragraph_format.line_spacing = 1.25

            # --- 渲染行内文本 (处理 **加粗**) ---
            # 逻辑：将文本按 ** 分割，偶数索引为普通，奇数索引为加粗
            # 例子: "这是 **重点** 内容" -> ["这是 ", "**重点**", " 内容"]
            # (split不保留分隔符，所以需要手动正则split保持捕获组，或者简单遍历)

            parts = re_bold.split(clean_text)
            # re.split 包含捕获组时：['这是 ', '**重点**', '重点', ' 内容']
            # 我们需要过滤掉那个内部捕获组 (group 2)，只保留 group 1 或者自己重写逻辑

            # 更简单的逻辑：手动遍历
            last_idx = 0
            for m in re_bold.finditer(clean_text):
                # 1. 添加前面的普通文本
                start, end = m.span()
                if start > last_idx:
                    text_chunk = clean_text[last_idx:start]
                    run = p.add_run(text_chunk)
                    self.set_font(run, self.SIZE_5, bold=is_bold_line)

                # 2. 添加加粗文本 (去掉 **)
                bold_content = m.group(2)  # 获取内部内容
                run = p.add_run(bold_content)
                self.set_font(run, self.SIZE_5, bold=True)  # 强制加粗

                last_idx = end

            # 3. 添加剩余文本
            if last_idx < len(clean_text):
                text_chunk = clean_text[last_idx:]
                run = p.add_run(text_chunk)
                self.set_font(run, self.SIZE_5, bold=is_bold_line)

    def _create_info_table(self, metadata):
        # 保持原有逻辑，确保引用正确
        table = self.doc.add_table(rows=4, cols=4)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False

        widths = [3.75, 4.75, 3.5, 5.5]
        for row in table.rows:
            row.height = Cm(1.0)
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            for idx, cell in enumerate(row.cells):
                cell.width = Cm(widths[idx])
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        def fill(r, c, text, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, merge_cols=None):
            cell = table.cell(r, c)
            if merge_cols:
                cell.merge(table.cell(r, c + merge_cols))
            if text is not None:
                p = cell.paragraphs[0]
                p.alignment = align
                run = p.runs[0] if p.runs else p.add_run(str(text))
                self.set_font(run, self.SIZE_5, bold=bold)
            return cell

        fill(0, 0, "课程名称")
        fill(0, 1, metadata.get('course_name', ''), merge_cols=2)
        fill(1, 0, "专业年级班级")
        fill(1, 1, metadata.get('class_name', ''), merge_cols=2)
        fill(2, 0, "考核形式")
        fill(2, 1, metadata.get('assessment_type', '考查'))
        fill(2, 2, "命题日期")
        fill(2, 3, metadata.get('date', ''))

        fill(3, 0, "命题教师")

        # 教师签名单元格
        c_teacher = table.cell(3, 1)
        p_t = c_teacher.paragraphs[0]
        p_t.alignment = WD_ALIGN_PARAGRAPH.LEFT
        tn = metadata.get('teacher_name', '')
        if tn:
            r = p_t.add_run(f"{tn}  ")
            self.set_font(r, self.SIZE_5, bold=True)
        if metadata.get('teacher_sig_path') and os.path.exists(metadata.get('teacher_sig_path')):
            try:
                p_t.add_run().add_picture(metadata.get('teacher_sig_path'), height=Cm(0.8))
            except:
                pass

        fill(3, 2, "系（教研室）\n主任审核签字")

        # 主任签名单元格
        c_head = table.cell(3, 3)
        p_h = c_head.paragraphs[0]
        p_h.alignment = WD_ALIGN_PARAGRAPH.LEFT
        hn = metadata.get('head_name', '')
        if hn:
            r = p_h.add_run(f"{hn}  ")
            self.set_font(r, self.SIZE_5, bold=True)
        if metadata.get('head_sig_path') and os.path.exists(metadata.get('head_sig_path')):
            try:
                p_h.add_run().add_picture(metadata.get('head_sig_path'), height=Cm(0.8))
            except:
                pass

    def _create_footer_notes(self):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.line_spacing = 1.0

        segments = [
            ("注：\n", False),
            ("1．课程名称必须与教学计划上的名称一致。\n", False),
            ("2. ", False),
            ("命题教师：", True),
            ("务必输入命题教师名字，打印纸质版后再手写签名；", False),
            ("系（教研室）主任审核签字：", True),
            ("须手写签名。\n", False),
            ("3. 该表文字部分均用五号宋体，使用A4纸双面打印。\n", False),
            ("4. 命题完成后将该表与命题计划表（电子版及纸质版）交到二级学院（部），并装入试卷袋存档。", False)
        ]

        for text, is_bold in segments:
            run = p.add_run(text)
            self.set_font(run, self.SIZE_5, bold=is_bold)

    def _create_footer(self, section):
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.add_page_number_field(p)