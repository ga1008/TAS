# export_core/syllabus_exporter.py
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from export_core.base_template import BaseExportTemplate
from export_core.word_exporter_base import BaseWordExporter


class SyllabusExporter(BaseExportTemplate, BaseWordExporter):
    ID = "syllabus_general"
    NAME = "通用课程教学大纲"
    DESCRIPTION = "适用于一般理论或实践课程的教学大纲标准模板。"

    # 前端 UI Schema：用于在导出页面生成表单
    # 利用 auto_fill_key 自动从 AI 提取的 metadata 中填充
    UI_SCHEMA = [
        {"name": "course_name", "label": "课程名称", "type": "text", "auto_fill_key": "course_name"},
        {"name": "course_code", "label": "课程编号", "type": "text", "auto_fill_key": "course_code"},

        # 组合字段：基本信息
        {"name": "basic_info", "label": "基本信息", "type": "group", "children": [
            {"name": "credits", "label": "学分", "type": "text", "width": "33%", "auto_fill_key": "credits"},
            {"name": "hours_total", "label": "总学时", "type": "text", "width": "33%",
             "auto_fill_key": "hours_info.total"},  # 注意：需前端或后端支持嵌套key读取
            {"name": "department", "label": "开课部门", "type": "text", "width": "33%", "auto_fill_key": "department"},
        ]},

        {"name": "course_category", "label": "课程类别", "type": "select",
         "options": ["通识必修课程", "通识选修课程", "专业基础课程", "专业核心课程", "专业限选课程", "专业任选课程"],
         "auto_fill_key": "course_category"},

        # 签名部分
        {"name": "drafter_name", "label": "制定人", "type": "text", "auto_fill_key": "authors.drafter"},
        {"name": "reviewer_name", "label": "审定人", "type": "text", "auto_fill_key": "authors.reviewer"},
    ]

    def __init__(self):
        BaseWordExporter.__init__(self)

    def generate(self, content, meta_info, form_data, output_path):
        # 1. 页面设置
        self.setup_page()

        # 2. 标题区
        p_title = self.doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 类似 《Python程序设计》课程教学大纲
        title_text = f"《{form_data.get('course_name', '未命名课程')}》课程教学大纲"
        self.set_font(p_title.add_run(title_text), self.SIZE_SMALL_2, bold=True, font_name=self.FONT_CN_BOLD)

        # 3. 课程编号与类别
        p_info = self.doc.add_paragraph()
        p_info.alignment = WD_ALIGN_PARAGRAPH.LEFT
        self.set_font(p_info.add_run(f"课程编号：{form_data.get('course_code', '')}"), self.SIZE_5)
        # 课程类别处理 (勾选框模拟)
        # 这里仅作简单的文本展示，实际可根据 form_data['course_category'] 动态画方框
        p_info.add_run("\n")
        self.set_font(p_info.add_run(f"课程类别：{form_data.get('course_category', '')}"), self.SIZE_5)

        # 4. 学时学分表 (示例)
        table = self.doc.add_table(rows=2, cols=4)
        table.style = 'Table Grid'
        # ... 填充表头和内容 (略，参考 GuangWaiStandardExporter) ...

        # 5. 正文内容 (Markdown -> Word)
        # 使用基类中的 clean_markdown 和结构化解析逻辑
        # 这里需要针对大纲的结构（一、二、三...）进行针对性渲染
        self._render_syllabus_body(content)

        # 6. 底部签字栏
        p_footer = self.doc.add_paragraph()
        p_footer.paragraph_format.space_before = Pt(24)
        p_footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        self.set_font(p_footer.add_run(f"制定人：{form_data.get('drafter_name', '')}    "), self.SIZE_4)
        self.set_font(p_footer.add_run(f"审定人：{form_data.get('reviewer_name', '')}"), self.SIZE_4)

        self.doc.save(output_path)
        return output_path

    def _render_syllabus_body(self, content):
        """
        专门渲染大纲正文，处理特定的层级结构
        """
        lines = content.split('\n')
        for line in lines:
            line = line.strip()
            if not line: continue

            p = self.doc.add_paragraph()

            # 简单判断标题级别
            if line.startswith('# '):
                # 忽略一级标题，因为上面已经生成了
                continue
            elif line.startswith('## ') or line.startswith('一、') or line.startswith('二、'):
                clean = line.replace('## ', '')
                self.set_font(p.add_run(clean), self.SIZE_4, bold=True, font_name=self.FONT_CN_BOLD)
            elif line.startswith('### ') or line.startswith('（一）') or line.startswith('（二）'):
                clean = line.replace('### ', '')
                p.paragraph_format.first_line_indent = Pt(21)
                self.set_font(p.add_run(clean), self.SIZE_5, bold=True)
            else:
                # 正文
                p.paragraph_format.first_line_indent = Pt(21)
                self.set_font(p.add_run(line), self.SIZE_5)
