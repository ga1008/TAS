# export_core/templates/plan_exporter.py

import os
import re
from datetime import datetime

from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

# [新增] 引入底层 XML 操作库，修复边距设置报错问题
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# 引入基类和工具类
from export_core.base_template import BaseExportTemplate
from export_core.word_exporter_base import BaseWordExporter


class AssessmentPlanExporter(BaseExportTemplate, BaseWordExporter):
    ID = "plan_guangwai"
    NAME = "广西外国语学院 - 考核计划表"
    DESCRIPTION = "适用于非笔试/机试等考核形式的计划表导出，包含双表格布局。"

    # === 1. 前端 UI Schema (表单定义) ===
    UI_SCHEMA = [
        {"name": "course_name", "label": "课程名称", "type": "text", "auto_fill_key": "course_name"},
        {"name": "class_name", "label": "专业年级班级", "type": "text", "auto_fill_key": "class_info"},

        {"name": "semester_info", "label": "学期信息", "type": "group", "children": [
            {"name": "year_start", "label": "起始年份", "type": "number", "width": "30%",
             "auto_fill_key": "academic_year_start"},
            {"name": "year_end", "label": "结束年份", "type": "number", "width": "30%",
             "auto_fill_key": "academic_year_end"},
            {"name": "semester", "label": "学期", "type": "select", "options": ["一", "二"], "width": "30%",
             "auto_fill_key": "semester"}
        ]},

        {"name": "assessment_type", "label": "考核类型", "type": "select", "options": ["考试", "考查"],
         "auto_fill_key": "assessment_type"},
        {"name": "assessment_note", "label": "考核说明", "type": "text", "placeholder": "（非笔试考核）",
         "auto_fill_key": "assessment_note"},

        {"name": "teacher_name", "label": "命题教师", "type": "text", "auto_fill_key": "teacher"},
        {"name": "teacher_sig", "label": "教师签名", "type": "signature_selector", "bind_to": "teacher_name"},

        {"name": "head_name", "label": "系主任", "type": "text", "auto_fill_key": "dept_head"},
        {"name": "head_sig", "label": "系主任签名", "type": "signature_selector", "bind_to": "head_name"},

        {"name": "date", "label": "命题日期", "type": "date", "auto_fill_key": "date"}
    ]

    def __init__(self):
        BaseWordExporter.__init__(self)

    # === [新增] 修复报错的辅助方法 ===
    def _set_cell_margins(self, cell, left_cm=0.19, right_cm=0.19):
        """
        通过底层 OXML 手动设置单元格边距 (修复 get_or_add_tcMar 报错)
        注意：w:type="dxa" 时，单位必须是 Twips (1cm ≈ 567 Twips)
        """
        tcPr = cell._element.get_or_add_tcPr()
        tcMar = tcPr.find(qn('w:tcMar'))
        if tcMar is None:
            tcMar = OxmlElement('w:tcMar')
            tcPr.append(tcMar)

        # 计算 Twips 值 (0.19cm * 567 ≈ 108)
        left_val = str(int(left_cm * 567))
        right_val = str(int(right_cm * 567))

        # 设置左边距
        node_left = tcMar.find(qn('w:left'))
        if node_left is None:
            node_left = OxmlElement('w:left')
            tcMar.append(node_left)
        node_left.set(qn('w:w'), left_val)
        node_left.set(qn('w:type'), 'dxa')

        # 设置右边距
        node_right = tcMar.find(qn('w:right'))
        if node_right is None:
            node_right = OxmlElement('w:right')
            tcMar.append(node_right)
        node_right.set(qn('w:w'), right_val)
        node_right.set(qn('w:type'), 'dxa')

    # === 2. 核心生成入口 ===
    def generate(self, content, meta_info, form_data, output_path):
        """
        生成 Word 文档的主入口
        """
        # 1. 页面基本设置 (A4, 边距均为 1.5cm)
        section = self.setup_page(1.5, 1.5, 1.5, 1.5, footer_distance=1.75)

        # 2. 大标题
        p_title = self.doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 宋体小二加粗
        self.set_font(p_title.add_run("广西外国语学院课程考核计划表"), self.SIZE_SMALL_2, bold=True,
                      font_name=self.FONT_CN_BOLD)

        # 3. 学年度信息 (宋体四号加粗)
        self._create_semester_header(form_data)

        # 4. 考核形式提示语 (宋体小四)
        note_text = form_data.get('assessment_note', '（非笔试考核）') or '（非笔试考核）'
        if not note_text.startswith("（"):
            note_text = f"（{note_text}）"
        p_note = self.doc.add_paragraph()
        p_note.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p_note.paragraph_format.line_spacing = 1.0
        p_note.paragraph_format.space_before = Pt(0)
        p_note.paragraph_format.space_after = Pt(0)

        self.set_font(p_note.add_run(note_text), self.SIZE_SMALL_4)

        # 5. 表格1：课程基本信息表
        self._create_info_table(form_data)

        # 增加一点间距
        self.doc.add_paragraph().paragraph_format.space_after = Pt(0)

        # 6. 表格2：考核计划正文表 (解析 Content 填充)
        self._create_plan_table(content)

        # 7. 底部备注
        self._create_footer_notes()

        # 8. 保存
        self.doc.save(output_path)
        return output_path

    # === 3. 内部组件实现 ===

    def _create_semester_header(self, data):
        """渲染学期信息：(20 25 - 20 26 学年度第 一 学期)"""
        y_start = str(data.get('year_start', ''))
        y_end = str(data.get('year_end', ''))
        semester = str(data.get('semester', ''))

        if not (y_start and y_end and semester):
            y_s_calc, y_e_calc, sem_calc = self.get_semester_info()
            y_start = y_start or y_s_calc
            y_end = y_end or y_e_calc
            semester = semester or sem_calc

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)

        self.set_font(p.add_run("（ "), self.SIZE_4, bold=True)

        def add_year_runs(year_str):
            if len(year_str) == 4:
                prefix = year_str[:2]  # 20
                suffix = year_str[2:]  # 25
            else:
                prefix = ""
                suffix = year_str

            self.set_font(p.add_run(prefix), self.SIZE_4, bold=True)
            # 后缀加空格和下划线
            r = p.add_run(f" {suffix} ")
            self.set_font(r, self.SIZE_4, bold=True, underline=True)

        add_year_runs(y_start)
        self.set_font(p.add_run(" — "), self.SIZE_4, bold=True)
        add_year_runs(y_end)

        self.set_font(p.add_run(" 学年度第 "), self.SIZE_4, bold=True)

        r_sem = p.add_run(f" {semester} ")
        self.set_font(r_sem, self.SIZE_4, bold=True, underline=True)

        self.set_font(p.add_run(" 学期）"), self.SIZE_4, bold=True)

    def _create_info_table(self, data):
        """创建表格1：课程信息表 (四行四列)"""
        table = self.doc.add_table(rows=4, cols=4)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False

        # 设置列宽
        # 第一列 4.64cm, 第二列 4.31cm, 第三列 4.25cm, 第四列 5.05cm
        widths = [4.64, 4.31, 4.25, 5.05]

        for row in table.rows:
            # 固定行高 1.1cm
            row.height = Cm(1.1)
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            for idx, cell in enumerate(row.cells):
                cell.width = Cm(widths[idx])
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                # [修复] 调用安全的边距设置方法
                self._set_cell_margins(cell, left_cm=0.19, right_cm=0.19)

        # 辅助填充函数
        def fill(r, c, text, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, merge_cols=None):
            cell = table.cell(r, c)
            if merge_cols:
                cell.merge(table.cell(r, c + merge_cols))

            # 获取合并后的段落
            if text is not None:
                # 清除可能存在的旧段落（合并导致）
                p = cell.paragraphs[0]
                p.alignment = align
                # 清除现有的 runs
                p.clear()
                run = p.add_run(str(text))
                # 全表五号加粗
                self.set_font(run, self.SIZE_5, bold=True)
            return cell

        # 第一行：课程名称 (合并 2,3,4 列)
        fill(0, 0, "课程名称")
        fill(0, 1, data.get('course_name', ''), merge_cols=2)

        # 第二行：专业年级班级 | 班级名 | 考核类型 | 考查/考试
        fill(1, 0, "专业年级班级")
        fill(1, 1, data.get('class_name', ''))
        fill(1, 2, "考核类型")

        # 动态生成考核类型勾选文本
        a_type = data.get('assessment_type', '考试')
        if a_type == '考查':
            check_text = "考查( √ ) / 考试(   )"
        else:
            check_text = "考查(   ) / 考试( √ )"
        fill(1, 3, check_text)

        # 第三行：命题教师 | 姓名+签名 | 审核签字 | 姓名+签名
        fill(2, 0, "命题教师")

        # 教师签名单元格 (左对齐)
        c_teacher = table.cell(2, 1)
        c_teacher.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p_t = c_teacher.paragraphs[0]
        p_t.clear()
        p_t.alignment = WD_ALIGN_PARAGRAPH.LEFT
        tn = data.get('teacher_name', '')
        if tn:
            r = p_t.add_run(f"{tn}  ")
            self.set_font(r, self.SIZE_5, bold=True)
        # 插入签名图
        if data.get('teacher_sig_path') and os.path.exists(data.get('teacher_sig_path')):
            try:
                r_img = p_t.add_run()
                r_img.add_picture(data.get('teacher_sig_path'), height=Cm(0.8))
            except:
                pass

        fill(2, 2, "系（教研室）\n主任审核签字")

        # 系主任签名单元格 (左对齐)
        c_head = table.cell(2, 3)
        c_head.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p_h = c_head.paragraphs[0]
        p_h.clear()
        p_h.alignment = WD_ALIGN_PARAGRAPH.LEFT
        hn = data.get('head_name', '')
        if hn:
            r = p_h.add_run(f"{hn}  ")
            self.set_font(r, self.SIZE_5, bold=True)
        if data.get('head_sig_path') and os.path.exists(data.get('head_sig_path')):
            try:
                r_img = p_h.add_run()
                r_img.add_picture(data.get('head_sig_path'), height=Cm(0.8))
            except:
                pass

        # 第四行：命题日期 (合并 2,3,4 列)
        fill(3, 0, "命题日期")
        date_val = data.get('date', '')
        if not date_val:
            now = datetime.now()
            date_val = f"{now.year}年{now.month}月{now.day}日"
        fill(3, 1, date_val, merge_cols=2)

    def _create_plan_table(self, content):
        """
        创建表格2：考核计划正文表
        需要解析 Markdown 表格内容填充。
        """
        # 1. 简单的 Markdown 表格解析器
        # 假设 content 是标准的 Markdown 表格
        rows_data = []
        lines = content.strip().split('\n')

        for line in lines:
            line = line.strip()
            # 跳过分隔行 |---|---|
            if set(line) <= {'|', '-', ' ', ':'}:
                continue
            if line.startswith('|') and line.endswith('|'):
                # 提取单元格内容
                cols = [c.strip() for c in line.split('|')[1:-1]]
                if len(cols) >= 3:
                    # 只取前三列，多余的合并到第三列或忽略
                    rows_data.append(cols[:3])

        # 如果解析不到数据，给一个默认空行
        if not rows_data:
            rows_data = [["", "", ""]]

        # 过滤掉表头 (如果 AI 生成了表头 "考核形式" 等，需要去掉，因为我们要自己画)
        if "考核形式" in rows_data[0][0]:
            rows_data = rows_data[1:]

        # 2. 创建 Word 表格
        # 行数 = 1 (表头) + 数据行数
        total_rows = 1 + len(rows_data)
        table = self.doc.add_table(rows=total_rows, cols=3)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False

        widths = [4.64, 10.11, 3.5]

        # 设置表头
        header_cells = table.rows[0].cells
        headers = ["考核形式", "考核技能/内容", "分 值"]
        table.rows[0].height = Cm(1.1)
        table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

        for idx, cell in enumerate(header_cells):
            cell.width = Cm(widths[idx])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            # [修复] 调用安全的边距设置方法
            self._set_cell_margins(cell, left_cm=0.19, right_cm=0.19)

            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(headers[idx])
            self.set_font(r, self.SIZE_5, bold=True)  # 五号加粗

        # 填充数据行
        for r_idx, row_data in enumerate(rows_data):
            # 表格行的索引需 +1 (跳过表头)
            word_row = table.rows[r_idx + 1]
            # 默认高度 1.1cm，自动拉伸
            word_row.height = Cm(1.1)
            # height_rule 默认为 AUTO (At Least)，适合自动换行拉高

            for c_idx, text in enumerate(row_data):
                cell = word_row.cells[c_idx]
                cell.width = Cm(widths[c_idx])
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                # [修复] 调用安全的边距设置方法
                self._set_cell_margins(cell, left_cm=0.19, right_cm=0.19)

                p = cell.paragraphs[0]
                # 对齐方式：第二列左对齐，其他居中
                if c_idx == 1:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # 处理换行符 <br>
                clean_text = text.replace("<br>", "\n").replace("<br/>", "\n")

                r = p.add_run(clean_text)
                self.set_font(r, self.SIZE_5, bold=False)  # 五号正常

    def _create_footer_notes(self):
        """生成底部备注信息"""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.line_spacing = 1.0  # 单倍行距
        p.paragraph_format.first_line_indent = 0  # 缩进0字符
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        note_content = """注：
1．课程名称必须与教学计划上的名称一致。
2．考核类型：考查、考试（按教学计划填写）。
3．命题教师：务必输入命题教师名字，打印纸质版后再手写签名；系（教研室）主任审核签字：须手写签名。
4．各专业根据教学大纲自行拟定考核形式、考核技能/内容、分值。
5. 该表文字部分均用五号宋体，使用A4纸双面打印。
6. 命题完成后将该表与评分细则（电子版及纸质版）交到二级学院（部），并装入试卷袋存档。"""

        # 逐行添加，保持格式统一
        for line in note_content.split('\n'):
            run = p.add_run(line + '\n')
            self.set_font(run, self.SIZE_5, bold=False)