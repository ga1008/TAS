import asyncio
import base64
import hashlib
import json
import mimetypes
import os
import re
import shutil
import threading
import time
import traceback
import uuid

import PyPDF2
import httpx
import pandas as pd
import patoolib
from docx import Document
from flask import Flask, render_template, request, redirect, url_for, send_file, jsonify, session, g

from ai_utils.ai_helper import call_ai_platform_chat
from ai_utils.volc_file_manager import VolcFileManager
from blueprints.admin import bp as admin_bp
from config import Config, BASE_CREATOR_PROMPT, STRICT_MODE_PROMPT, LOOSE_MODE_PROMPT, EXAMPLE_PROMPT
from export_core.doc_config import DocumentTypeConfig
from export_core.manager import TemplateManager
# 确保 extensions.py 中已正确初始化 db = Database()
from extensions import db
from grading_core.direct_grader_template import DIRECT_GRADER_TEMPLATE
from grading_core.factory import GraderFactory

app = Flask(__name__)
app.config.from_object(Config)

# 注册蓝图
app.register_blueprint(admin_bp)

# 确保基础目录存在
for d in [
    app.config['UPLOAD_FOLDER'],
    app.config['WORKSPACE_FOLDER'],
    app.config['GRADERS_DIR'],
    app.config['TRASH_DIR'],
    app.config['FILE_REPO_FOLDER'],
    app.config['TEMPLATE_DIR'],
    app.config['SIGNATURES_FOLDER']
]:
    if not os.path.exists(d): os.makedirs(d)

TemplateManager.load_templates(app.config['TEMPLATE_DIR'])

# 启动时加载评分策略
try:
    GraderFactory.load_graders()
except Exception as e:
    print(f"Warning: Failed to load graders on startup: {e}")


@app.before_request
def load_logged_in_user():
    user_info = session.get('user')
    if user_info:
        conn = db.get_connection()
        user = conn.execute("SELECT * FROM users WHERE id=?", (user_info['id'],)).fetchone()
        g.user = dict(user) if user else None
    else:
        g.user = None


# === 工具函数 ===

# === 文件处理工具 ===
def calculate_file_hash(file_stream):
    """计算文件的 SHA256 哈希值"""
    sha256 = hashlib.sha256()
    # 读取指针归零
    file_stream.seek(0)
    while chunk := file_stream.read(8192):
        sha256.update(chunk)
    # 计算完后将指针归零，方便后续保存
    file_stream.seek(0)
    return sha256.hexdigest()


def get_real_workspace_path(class_id):
    """
    不管数据库里存的路径是什么（可能是旧的Windows路径），
    强制使用当前配置的 WORKSPACE_FOLDER 拼接 class_id 生成正确路径。
    """
    # 确保 WORKSPACE_FOLDER 存在
    if not os.path.exists(app.config['WORKSPACE_FOLDER']):
        os.makedirs(app.config['WORKSPACE_FOLDER'])

    ws_path = os.path.join(app.config['WORKSPACE_FOLDER'], str(class_id))

    # 确保班级目录存在
    if not os.path.exists(ws_path):
        os.makedirs(ws_path)

    return ws_path


def _grade_single_student_internal(class_id, student_id):
    """
    内部核心逻辑：只批改一个学生
    """
    # 1. 获取基本信息
    cls_info = db.get_class_by_id(class_id)
    if not cls_info:
        return False, "班级不存在", None

    # 务必使用 get_real_workspace_path (如果你在之前步骤已添加该函数)
    # 如果没有添加，请确保这里使用正确的路径逻辑
    if 'get_real_workspace_path' in globals():
        workspace_path = get_real_workspace_path(class_id)
    else:
        # 兼容旧逻辑
        workspace_path = cls_info['workspace_path']

    raw_dir = os.path.join(workspace_path, 'raw_zips')
    extract_base = os.path.join(workspace_path, 'extracted')

    # 2. 获取学生信息
    conn = db.get_connection()
    student = conn.execute("SELECT * FROM students WHERE class_id=? AND student_id=?",
                           (class_id, student_id)).fetchone()
    if not student:
        return False, "找不到学生记录", None

    name = student['name']

    # 3. 查找对应的压缩包
    if not os.path.exists(raw_dir):
        return False, "未上传任何文件", None

    uploaded_files = os.listdir(raw_dir)
    matched_file = None
    # 文件名匹配逻辑：学号 或 姓名
    for f in uploaded_files:
        if str(student_id) in f or name in f:
            matched_file = f
            break

    if not matched_file:
        db.save_grade_error(student_id, class_id, "未找到提交文件", "")
        return False, "未找到提交文件 (No Submission)", None

    # 4. 准备解压目录
    student_extract_dir = os.path.join(extract_base, str(student_id))
    if os.path.exists(student_extract_dir):
        try:
            shutil.rmtree(student_extract_dir)
        except:
            pass  # 忽略删除错误
    os.makedirs(student_extract_dir, exist_ok=True)

    archive_path = os.path.join(raw_dir, matched_file)

    try:
        # 5. 加载评分策略
        grader = GraderFactory.get_grader(cls_info['strategy'])
        if not grader:
            return False, f"策略 {cls_info['strategy']} 加载失败", matched_file

        # 6. 解压
        try:
            patoolib.extract_archive(archive_path, outdir=student_extract_dir, verbosity=-1)
        except Exception as e:
            # 针对 Linux 环境下缺少 rar 支持的特殊提示
            if "rar" in matched_file.lower() and "patool" in str(e):
                raise Exception("解压失败(服务器不支持RAR格式，请上传ZIP)")
            raise e

        # 7. === 核心：调用 AI 批改 ===
        # 这一步最耗时 (30s - 60s)
        result = grader.grade(student_extract_dir, {"sid": str(student_id), "name": name})

        # 8. 保存结果
        status = "PASS" if result.is_pass else "FAIL"
        db.save_grade(str(student_id), class_id, result.total_score, result.get_details_json(),
                      result.get_deduct_str(), status, matched_file)

        return True, "批改完成", {
            "total_score": result.total_score,
            "status": status,
            "filename": matched_file,
            "deduct": result.get_deduct_str(),
            "details": result.sub_scores
        }

    except Exception as e:
        err_msg = f"系统异常: {str(e)}"
        print(f"[Grading Error] {student_id}: {e}")
        # traceback.print_exc()
        db.save_grade_error(str(student_id), class_id, err_msg, matched_file)
        return False, err_msg, matched_file


# === [新增] 路径修复工具函数 ===
def get_corrected_path(db_path, folder_path):
    """
    智能修复路径：
    1. 如果数据库里的路径在当前服务器存在，直接用。
    2. 如果不存在（比如是从 Windows 迁移到 Linux），则提取文件名，拼接到当前的 folder_path 下再找一次。
    """
    if not db_path:
        return None

    # 情况A: 路径有效 (本地创建或路径一致)
    if os.path.exists(db_path):
        return db_path

    # 情况B: 路径无效 (跨平台迁移导致)
    # 提取文件名 (同时兼容 Windows 的 \ 和 Linux 的 /)
    filename = os.path.basename(db_path.replace('\\', '/'))
    corrected_path = os.path.join(folder_path, filename)

    if os.path.exists(corrected_path):
        return corrected_path

    return None


# app.py

# 引入新写的转换工具
from utils.file_converter import convert_to_pdf
from ai_utils.volc_file_manager import VolcFileManager

# app.py 部分代码替换

# 引入转换工具 (确保文件存在)
from utils.file_converter import convert_to_pdf


def smart_parse_with_metadata(file_id, doc_category_hint="exam"):
    """
    智能解析文件 V3 (视觉多模态版)：
    1. [新增] 尝试将 Doc/Docx 转为 PDF。
    2. [新增] 上传 PDF 到火山引擎 Files API。
    3. [新增] 使用 Vision 模型进行"视觉阅读"，完美保留表格结构。
    4. [兜底] 如果视觉失败，回退到纯文本提取。
    """
    import json
    import asyncio
    import re
    from extensions import db
    from ai_utils.ai_helper import call_ai_platform_chat

    # 1. 获取文件记录
    record = db.get_file_by_id(file_id)
    if not record:
        return False, "文件记录不存在", {}

    physical_path = record['physical_path']
    original_ext = os.path.splitext(physical_path)[1].lower()

    parsed_content = ""
    meta = {}
    used_vision = False
    error_log = []

    # === 策略 A: 视觉解析 (优先) ===
    # 只有当配置了 Vision 模型，且文件是文档类型时尝试
    vision_config = db.get_best_ai_config("vision")

    if vision_config and original_ext in ['.docx', '.doc', '.pdf']:
        print(f"[SmartParse] Attempting Vision Mode for {record['original_name']}...")
        try:
            # A1. 格式转换
            target_file_path = physical_path
            if original_ext != '.pdf':
                converted_pdf = convert_to_pdf(physical_path)
                if converted_pdf and os.path.exists(converted_pdf):
                    target_file_path = converted_pdf
                else:
                    raise Exception("PDF conversion failed")

            # A2. 上传到火山引擎
            uploader = VolcFileManager(api_key=vision_config['api_key'], base_url=vision_config.get('base_url'))
            remote_file_id = uploader.upload_file(target_file_path)

            if remote_file_id:
                print(f"[SmartParse] Uploaded to VolcEngine: {remote_file_id}")

                # A3. 构造视觉 Prompt
                type_specific_instruction = DocumentTypeConfig.get_prompt_by_type(doc_category_hint)

                # 增强视觉引导
                user_prompt = f"""
                {type_specific_instruction}

                【视觉特别指令】
                1. 请利用视觉能力识别文档中的布局。
                2. **特别注意勾选框**：文档中可能有 '□', '☑', 'R', '√' 等符号。请务必识别哪个选项被勾选了，并提取该选项的文字到 metadata 中。
                3. **表格识别**：请保持表格的行列对应关系，尤其是'考核计划表'的三列对应。
                4. 若有图片（如知识图谱），请尝试描述或用 Mermaid 代码表示。
                """

                # A4. 调用 AI (ai_helper 会自动切换到 Responses API)
                ai_response = asyncio.run(call_ai_platform_chat(
                    system_prompt="你是高校教学资料结构化专家。请基于视觉识别文档内容。",
                    messages=[{
                        "role": "user",
                        "content": user_prompt,
                        "file_ids": [remote_file_id]  # 传入 ID
                    }],
                    platform_config=vision_config
                ))

                # 简单的有效性检查
                if ai_response and len(ai_response) > 50 and "[PARSE_ERROR]" not in ai_response:
                    used_vision = True
                    raw_response_text = ai_response
                    print("[SmartParse] Vision Mode Success.")
                else:
                    error_log.append("Vision model returned empty or error.")
            else:
                error_log.append("Failed to upload file to VolcEngine.")

        except Exception as e:
            print(f"[SmartParse] Vision Mode Error: {e}")
            error_log.append(str(e))

    # === 策略 B: 纯文本提取 (兜底) ===
    if not used_vision:
        print(f"[SmartParse] Fallback to Text Mode. (Errors: {'; '.join(error_log)})")
        parse_success, raw_text = extract_text_from_file(physical_path)

        if not parse_success or not raw_text:
            return False, f"基础解析失败: {raw_text}", {}

        input_content = raw_text[:50000]
        type_specific_instruction = DocumentTypeConfig.get_prompt_by_type(doc_category_hint)

        user_prompt = f"""
            {type_specific_instruction}
            请务必严格遵守以下输出要求：
            1. **必须仅返回一个标准的 JSON 对象**。
            2. JSON 根对象包含 "content" 和 "metadata"。

            待处理内容：
            {input_content}
        """

        standard_config = db.get_best_ai_config("thinking") or db.get_best_ai_config("standard")
        if not standard_config:
            return False, "未配置 AI 模型", {}

        raw_response_text = asyncio.run(call_ai_platform_chat(
            system_prompt="你是高校教学资料结构化专家。",
            messages=[{"role": "user", "content": user_prompt}],
            platform_config=standard_config
        ))

    # === 通用：JSON 清洗与入库 ===
    try:
        cleaned_json = raw_response_text.strip()
        # Regex 提取 JSON
        json_match = re.search(r'```(?:json)?\s*(\{.*?\})\s*```', cleaned_json, re.DOTALL)
        if json_match:
            cleaned_json = json_match.group(1)
        else:
            start = cleaned_json.find('{')
            end = cleaned_json.rfind('}')
            if start != -1 and end != -1:
                cleaned_json = cleaned_json[start:end + 1]

        data = json.loads(cleaned_json)

        if not isinstance(data, dict): raise ValueError("Not a dict")

        parsed_content = data.get("content", "")
        meta = data.get("metadata", {})

        if not parsed_content: raise ValueError("No content")

        # 提取字段入库
        academic_year = meta.get("academic_year")
        semester = str(meta.get("semester", ""))
        course_name = meta.get("course_name")
        cohort_tag = meta.get("cohort_tag")

        conn = db.get_connection()
        conn.execute('''
                     UPDATE file_assets
                     SET parsed_content=?,
                         meta_info=?,
                         doc_category=?,
                         academic_year=?,
                         semester=?,
                         course_name=?,
                         cohort_tag=?
                     WHERE id = ?
                     ''', (
                         parsed_content,
                         json.dumps(meta, ensure_ascii=False),
                         doc_category_hint,
                         academic_year, semester, course_name, cohort_tag,
                         file_id
                     ))
        conn.commit()

        return True, parsed_content, meta

    except Exception as e:
        print(f"[SmartParse] Parse/Save Error: {e}")
        # 如果是 Vision 模式失败，可能返回特定错误
        if used_vision:
            return False, f"视觉解析结果处理失败: {str(e)}", {}
        return False, f"解析失败: {str(e)}", {}


def handle_file_upload_or_reuse(file_obj, existing_file_id, user_id):
    """
    处理文件上传逻辑：
    1. 如果传入了 existing_file_id，直接查库返回路径。
    2. 如果传入了 file_obj，计算哈希 -> 查重 -> 保存/复用 -> 存库 -> 返回路径。
    """
    # 情况 A: 用户选择了已有文件
    if existing_file_id:
        record = db.get_file_by_id(existing_file_id)
        # 使用 os.path.normpath 规范化路径，无需手动处理字符串
        if record:
            return record['physical_path'], record['original_name']
        else:
            raise Exception("选中的文件记录不存在或物理文件丢失")

    # 情况 B: 用户上传了新文件
    if file_obj and file_obj.filename:
        # 1. 计算哈希
        f_hash = calculate_file_hash(file_obj)

        # 2. 查重
        existing_record = db.get_file_by_hash(f_hash)

        if existing_record:
            if os.path.exists(existing_record['physical_path']):
                return existing_record['physical_path'], file_obj.filename
            else:
                pass  # 数据库有记录但文件没了，继续向下执行覆盖保存

        # 3. 保存新文件
        ext = os.path.splitext(file_obj.filename)[1]
        save_filename = f"{f_hash}{ext}"
        physical_path = os.path.join(app.config['FILE_REPO_FOLDER'], save_filename)

        # === 【关键修改】 ===
        # 1. 删除了 physical_path.replace('\\', '\\\\')
        # 2. 使用 normpath 确保路径在不同系统下都标准
        physical_path = os.path.normpath(physical_path)
        file_obj.save(physical_path)
        file_size = os.path.getsize(physical_path)

        # 4. 写入数据库
        db.save_file_asset(f_hash, file_obj.filename, file_size, physical_path, user_id)

        return physical_path, file_obj.filename

    return None, None


def sanitize_filename(name, fallback="未命名文档"):
    clean = re.sub(r'[\\/*?:"<>|]', '', name or '').strip()
    return clean or fallback


def generate_title_from_content(content, doc_type=None):
    lines = [line.strip() for line in content.splitlines() if line.strip()]
    if lines:
        candidate = re.sub(r'^#+\s*', '', lines[0]).strip()
        if candidate:
            return sanitize_filename(candidate[:40])

    prefix = "试卷" if doc_type == "exam" else "评分标准"
    timestamp = time.strftime("%Y%m%d_%H%M")
    return f"{prefix}_{timestamp}"


def generate_title_with_ai(content, doc_type=None):
    doc_label = "试卷" if doc_type == "exam" else "评分标准"
    standard_config = db.get_best_ai_config("standard") or db.get_best_ai_config("thinking")
    if not standard_config:
        return generate_title_from_content(content, doc_type)

    prompt = (
        f"请为以下{doc_label}内容生成一个简洁准确的中文标题，20字以内。\n"
        "只返回标题文字，不要附加说明或标点。\n"
        f"内容：\n{content}\n"
    )
    try:
        ai_title = asyncio.run(call_ai_platform_chat(
            system_prompt="你是教学资料命名助手，擅长总结试卷或评分细则的标题。",
            messages=[{"role": "user", "content": prompt}],
            platform_config=standard_config
        ))
        ai_title = sanitize_filename((ai_title or "").strip())
        return ai_title or generate_title_from_content(content, doc_type)
    except Exception:
        return generate_title_from_content(content, doc_type)


def create_text_asset(content, title, user_id, ext=".md"):
    encoded = content.encode('utf-8')
    file_hash = hashlib.sha256(encoded).hexdigest()
    safe_title = sanitize_filename(title)
    filename = f"{safe_title}{ext}"
    physical_path = os.path.normpath(os.path.join(app.config['FILE_REPO_FOLDER'], f"{file_hash}{ext}"))

    if not os.path.exists(physical_path):
        with open(physical_path, 'wb') as f:
            f.write(encoded)

    file_id = db.save_file_asset(file_hash, filename, len(encoded), physical_path, user_id)
    if file_id:
        db.update_file_parsed_content(file_id, content)
    return file_id, filename


def extract_title_and_content(ai_text, doc_type=None):
    """
    解析 AI 返回的 JSON 或 文本。
    返回: (title, content, metadata_dict)
    """
    cleaned = ai_text.strip()
    meta = {}
    content = ""
    title = ""

    # 1. 尝试 Regex 提取 JSON
    json_match = re.search(r'```(?:json)?\s*(\{.*?\})\s*```', cleaned, re.DOTALL)
    if json_match:
        cleaned = json_match.group(1)

    try:
        # 尝试寻找 JSON 边界
        start = cleaned.find('{')
        end = cleaned.rfind('}')
        if start != -1 and end != -1:
            potential_json = cleaned[start:end + 1]
            payload = json.loads(potential_json)

            # 兼容新旧格式
            content = payload.get("content") or payload.get("body") or ""
            meta = payload.get("metadata") or {}

            # 尝试从 meta 拼凑标题
            if meta.get("course_name"):
                suffix = DocumentTypeConfig.TYPES.get(doc_type, "文档")
                title = f"{meta['course_name']} {suffix}"
            elif payload.get("title"):
                title = payload.get("title")

    except Exception:
        # JSON 解析失败，视为纯文本
        content = ai_text

    if not content: content = ai_text  # 兜底
    if not title: title = generate_title_from_content(content, doc_type)

    return sanitize_filename(title), content


def is_content_garbage(text):
    """
    检测文本是否为乱码/二进制内容
    特征：包含空字节、过多的控制字符、或常见的二进制头标识
    """
    if not text:
        return False

    # 1. 致命特征：包含空字节 (Null Byte) -> 必定是二进制乱码
    if '\x00' in text:
        return True

    # 2. 统计控制字符密度 (排除换行、回车、制表符)
    # 普通文本中，控制字符极少；二进制文件中，控制字符很密集
    total_len = len(text)
    if total_len == 0: return False

    # ASCII 0-8, 11-12, 14-31 是不可打印的控制字符
    control_chars = [c for c in text if 0 <= ord(c) < 32 and c not in ('\t', '\n', '\r')]
    control_ratio = len(control_chars) / total_len

    # 如果控制字符占比超过 5%，或者出现了常见的二进制头 (如 bjbj, PNG)，则判定为乱码
    # 你的例子中大约有 20%-30% 是控制字符
    if total_len > 50 and control_ratio > 0.05:
        return True

    # 3. 特征字检测 (针对旧版 Office 文档)
    if "bjbj" in text and control_ratio > 0.01:
        return True

    return False


def extract_text_from_file(file_path):
    """
    从不同格式文件中提取文本 (V2.2 增强健壮性版)
    """
    ext = os.path.splitext(file_path)[1].lower()
    text = ""

    # 错误净化内部函数
    def sanitize_error(e):
        msg = str(e)
        if os.path.sep in msg or "Users" in msg:
            print("[Sanitize Error] 文件格式解析错误:", msg)
            return "文件格式解析错误 (详情已隐藏)"
        return msg

    try:
        # 1. PDF
        if ext == '.pdf':
            reader = PyPDF2.PdfReader(file_path)
            for page in reader.pages:
                extracted = page.extract_text()
                if extracted: text += extracted + "\n"

        # 2. Word (.docx)
        elif ext == '.docx':
            doc = Document(file_path)
            for para in doc.paragraphs:
                text += para.text + "\n"

        # 3. 旧版 Word (.doc) 或 伪装格式
        elif ext == '.doc':
            # 尝试 A: 当做 docx
            try:
                doc = Document(file_path)
                for para in doc.paragraphs:
                    text += para.text + "\n"
            except:
                # 尝试 B: 当做 HTML/XML/MHTML (很多系统导出的 doc 其实是网页)
                try:
                    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                        content = f.read()
                        # 简单的 HTML 标签剥离
                        if '<html' in content.lower() or '<?xml' in content.lower():
                            text = re.sub(r'<[^>]+>', '\n', content)
                        else:
                            text = content
                except Exception as sub_e:
                    print(f"[Extract .doc fallback error] {sub_e}")
                    return False, "解析失败: 无法识别的 .doc 格式，建议另存为 .docx 或 .txt 后上传。"

        # 4. 纯文本/代码 (.txt, .md, .py, etc.)
        elif ext in ['.txt', '.py', '.java', '.c', '.cpp', '.md', '.json', '.html', '.css', '.js', '.sql', '.sh',
                     '.yaml', '.yml']:
            # 增强编码检测
            encodings = ['utf-8', 'gb18030', 'gbk', 'big5', 'latin1']
            for enc in encodings:
                try:
                    with open(file_path, 'r', encoding=enc) as f:
                        text = f.read()
                    break
                except UnicodeDecodeError:
                    continue

            # 如果依然失败，尝试二进制忽略错误读取
            if not text:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    text = f.read()

        else:
            # 未知格式，尝试纯文本读取
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()

    except Exception as e:
        print(f"[Extract Error] {file_path}: {e}")
        return False, f"解析失败: {sanitize_error(e)}"

    if is_content_garbage(text):
        print(f"[Garbage Detected] File {file_path} contains binary garbage.")
        # 返回特定的错误前缀，让上层逻辑识别
        return False, "解析失败: 文件似乎是二进制格式或已损坏（出现乱码）。请手动复制文本内容粘贴，或将文件转换为 .docx/.pdf 格式。"

    # 后处理：清理空字节和过多空白
    if text:
        text = text.replace('\x00', '')  # 去除空字节
        if not text.strip():
            return False, "解析结果为空 (文件可能为空或全为图片)"
        return True, text

    return False, "解析失败: 未能提取到有效文本"


# === 核心业务逻辑 ===

def normalize_document_with_ai(content, doc_type=None):
    doc_label = "试卷" if doc_type == "exam" else "评分标准"
    standard_config = db.get_best_ai_config("standard") or db.get_best_ai_config("thinking")
    if not standard_config:
        return None

    prompt = (
        f"请对以下{doc_label}内容进行规范化整理，仅保留试题或评分细则等核心内容。\n"
        "去除无用的表头、页眉、页脚、目录、页码等杂讯。\n"
        "输出 Markdown，并且不要遗漏任何主要信息。\n\n"
        f"{content}"
    )
    try:
        ai_text = asyncio.run(call_ai_platform_chat(
            system_prompt="你是文档规范化整理助手，擅长清理试卷与评分标准内容。",
            messages=[{"role": "user", "content": prompt}],
            platform_config=standard_config
        ))
        return ai_text.strip() if ai_text else None
    except Exception:
        return None


def smart_parse_file_content(file_id, doc_type=None):
    """
    智能解析文件内容 (V2.1 增强版)：
    1. 缓存优先。
    2. 尝试 AI Vision/Text 解析。
    3. 增加结果校验：防止 AI 返回 "无法解析" 的道歉文案。
    4. 自动降级：AI 失败时，强制使用本地 Python 库提取纯文本。
    """
    record = db.get_file_by_id(file_id)
    if not record:
        # raise Exception("文件记录不存在")
        return False, "解析失败: 文件记录不存在"

    # 1. 命中缓存
    if record.get('parsed_content'):
        print(f"[SmartParse] Cache hit for file {file_id}")
        return True, record['parsed_content'], False

    physical_path = record['physical_path']
    ext = os.path.splitext(physical_path)[1].lower()

    print(f"[SmartParse] Parsing file {file_id} ...")

    # === 增强 System Prompt ===
    doc_label = "试卷" if doc_type == "exam" else "评分标准"
    system_prompt = f"""
        你是一个专业的教学资料结构化专家。请读取用户上传的【{doc_label}】文件，将其整理为清晰、规范的 Markdown 格式。

        【核心原则】
        1. **完整性**：严禁遗漏任何细节！所有的题目描述、选项、分值、评分细则的扣分点、逻辑判断条件必须全部保留。
        2. **结构化**：
           - 使用一级标题 `#` 表示文档标题。
           - 使用二级标题 `##` 区分大题或主要模块（如“一、选择题”、“二、编程题”）。
           - 对于试题，请清晰标记题号（如 `**1.**`）。
           - 对于评分标准，请使用列表展示每一条得分点和扣分点。
        3. **易读性**：
           - 代码段请使用 ```code ... ``` 包裹。
           - 关键分值请使用加粗（如 **(5分)**）。
        4. **纯净输出**：直接输出 Markdown 内容，不要包含 "好的，这是解析结果" 等废话。
        5. **异常处理**：如果文件无法识别或损坏，仅输出：[PARSE_ERROR]
        """

    parsed_text = None
    parse_success = False
    used_fallback = False

    try:
        # 尝试获取 Vision 模型 (优先) 或 Standard 模型
        vision_config = db.get_best_ai_config("vision")

        # 定义支持 Vision 上传的类型
        is_media = ext in ['.pdf', '.jpg', '.jpeg', '.png', '.mp4', '.bmp', '.webp', '.tiff', '.docx', '.doc', '.txt',
                           '.py', '.java', '.c', '.md', '.json', '.html', '.js']

        # --- 分支 A: 调用 Vision 模型 ---
        if is_media and vision_config and vision_config.get('api_key'):
            uploader = VolcFileManager(api_key=vision_config['api_key'], base_url=vision_config.get('base_url'))
            remote_file_id = uploader.upload_file(physical_path)
            if remote_file_id:
                messages = [{
                    "role": "user",
                    "content": "请解析这份文档的全部内容。如果无法解析，请返回 [PARSE_ERROR]",
                    "file_ids": [remote_file_id]
                }]
                parsed_text = asyncio.run(call_ai_platform_chat(
                    system_prompt=system_prompt,
                    messages=messages,
                    platform_config=vision_config
                ))

        # --- 分支 B: 调用 Standard 模型 (处理文本/代码文件) ---
        elif ext in ['.txt', '.py', '.java', '.c', '.md', '.json', '.html', '.js']:
            parse_success_extract, raw_text = extract_text_from_file(physical_path)
            if parse_success_extract and len(raw_text) < 50000:
                standard_config = db.get_best_ai_config("standard") or db.get_best_ai_config("thinking")
                if standard_config:
                    parsed_text = asyncio.run(call_ai_platform_chat(
                        system_prompt=system_prompt,
                        messages=[{"role": "user", "content": f"请整理以下文本：\n\n{raw_text}"}],
                        platform_config=standard_config
                    ))

        # === 结果校验逻辑 ===
        failure_keywords = ["[PARSE_ERROR]", "无法直接解析", "provide the correct file"]
        if parsed_text:
            is_failure = any(k in parsed_text for k in failure_keywords)
            if is_failure or (len(parsed_text) < 50 and "无法" in parsed_text):
                parse_success = False
            else:
                parse_success = True

    except Exception as e:
        print(f"[SmartParse] AI Exception: {e}")
        parse_success = False

    # === 自动降级处理 (Fallback) ===
    if not parse_success:
        print("[SmartParse] Switching to Local Python Fallback Extraction...")
        # 强制使用本地逻辑提取文本
        parse_success, fallback_text = extract_text_from_file(physical_path)
        if not parse_success:
            return False, fallback_text  # 直接返回错误信息

        if fallback_text and len(fallback_text.strip()) > 0:
            print(f"### [系统提示：AI解析失败，已切换为纯文本提取模式]\n\n{fallback_text[:100]}")
        else:
            print(f"[SmartParse] [解析失败]\n无法从文件中提取有效文本，请检查文件是否损坏或已加密。")

        parsed_text = fallback_text
        used_fallback = True

        normalized_text = normalize_document_with_ai(fallback_text, doc_type)
        if normalized_text:
            parsed_text = normalized_text

    # 存入数据库
    if parsed_text:
        db.update_file_parsed_content(file_id, parsed_text)

    return parse_success, parsed_text, used_fallback


def ai_generation_worker(task_id, exam_text, standard_text, strictness, extra_desc, max_score):
    """
    后台线程：组装Prompt -> 调用 AI 微服务 -> 接收代码 -> 保存并热重载
    """

    def update_status(status, log, grader_id=None):
        try:
            db.update_ai_task(task_id, status=status, log_info=log, grader_id=grader_id)
        except Exception as e:
            print(f"Error updating task status: {e}")

    try:
        update_status("processing", "正在根据严格度组装 Prompt...")

        # 1. 动态组装 Prompt
        prompt_parts = [BASE_CREATOR_PROMPT]

        # 注入严格度要求
        strict_label = "标准模式"
        if strictness == 'strict':
            prompt_parts.append(STRICT_MODE_PROMPT)
            strict_label = "严格模式"
        elif strictness == 'loose':
            prompt_parts.append(LOOSE_MODE_PROMPT)
            strict_label = "宽松模式"
        else:
            prompt_parts.append(
                "### 3. 评分风格要求：【标准模式】\n请在严格匹配和容错之间保持平衡。对于关键错误扣分，对于格式小错误可以忽略。")

        prompt_parts[0] = prompt_parts[0].replace("{strictness_label}", strict_label)

        # 注入总分限制
        score_instruction = f"""
                ### 4. 分数控制 (Critical)
                - **总分限制**: 本次考试的满分必须严格等于 **{max_score}分**。
                - 请根据评分细则分配分数，确保 `self.res.total_score` 的计算逻辑最终汇总不超过也不低于 {max_score}。
                """
        prompt_parts.append(score_instruction)

        # 注入用户额外描述
        if extra_desc and extra_desc.strip():
            prompt_parts.append(
                f"\n### 5. 用户额外特殊指令 (优先级最高)\n用户特别强调：{extra_desc}\n请务必在代码中实现上述逻辑。")

        # 注入素材
        input_material = f"""
        ### 6. 输入素材
        --- [试卷内容开始] ---
        {exam_text}
        --- [试卷内容结束] ---

        --- [评分标准开始] ---
        {standard_text}
        --- [评分标准结束] ---
        """
        prompt_parts.append(input_material)

        prompt_parts.append(EXAMPLE_PROMPT)

        final_prompt = "\n".join(prompt_parts)

        # 2. 组装 Payload
        payload = {
            "system_prompt": "你是一名资深的 Python 自动化测试工程师。",
            "messages": [],
            "new_message": final_prompt,
            "model_capability": "thinking"
        }

        # 3. 发送请求给 AI 助手微服务
        update_status("processing", "已发送至 AI 助手，正在进行深度思考与代码生成，预计需要5分钟...")

        try:
            response = httpx.post(
                app.config['AI_ASSISTANT_CHAT_ENDPOINT'],
                json=payload,
                timeout=600.0
            )
        except httpx.ConnectError:
            raise Exception("无法连接到 AI 助手服务，请检查 ai_assistant.py 是否已启动 (端口9011)")

        if response.status_code != 200:
            raise Exception(f"AI 助手返回错误: {response.status_code} - {response.text}")

        ai_resp_json = response.json()
        ai_content = ai_resp_json.get("response_text", "")

        # 4. 解析返回的 Python 代码
        code_match = re.search(r'```python(.*?)```', ai_content, re.DOTALL)
        if not code_match:
            if "class " in ai_content and "BaseGrader" in ai_content:
                code_content = ai_content
            else:
                raise Exception("AI 返回内容中未找到有效的 Python 代码块")
        else:
            code_content = code_match.group(1).strip()

        # 5. 提取 ID 并保存
        id_match = re.search(r'ID\s*=\s*["\']([^"\']+)["\']', code_content)
        if not id_match:
            raise Exception("AI 生成的代码中缺少 ID 属性定义")

        grader_id = id_match.group(1)
        if not re.match(r'^[a-zA-Z0-9_]+$', grader_id):
            raise Exception(f"生成的 ID '{grader_id}' 格式不合法")

        graders_dir = os.path.join(os.path.dirname(__file__), 'grading_core', 'graders')
        if not os.path.exists(graders_dir):
            os.makedirs(graders_dir)

        filename = f"{grader_id}.py"
        save_path = os.path.join(graders_dir, filename)

        with open(save_path, 'w', encoding='utf-8') as f:
            f.write(code_content)
            f.flush()

        # 6. 热重载策略
        try:
            GraderFactory._loaded = False
            GraderFactory.load_graders()

            if grader_id in GraderFactory._graders:
                update_status("success", "核心构建完成，已自动挂载。", grader_id=grader_id)
            else:
                update_status("failed", "代码已生成但加载验证失败，请检查语法。", grader_id=grader_id)

        except Exception as e:
            update_status("failed", f"热重载失败: {str(e)}", grader_id=grader_id)

    except Exception as e:
        traceback.print_exc()
        update_status("failed", f"执行失败: {str(e)}")


@app.template_filter('split')
def split_filter(s, delimiter=None):
    if not s:
        return []
    return s.split(delimiter)


@app.template_filter('from_json')
def from_json_filter(s):
    if not s: return []
    try:
        return json.loads(s)
    except:
        return []


# === 路由定义 ===

@app.before_request
def auth_middleware():
    user_info = session.get('user')
    g.user = user_info if user_info else None
    if request.endpoint in ['login', 'logout', 'static', 'index']:
        return None
    if request.path.startswith('/admin'):
        return None
    if not g.user:
        return redirect(url_for('login'))
    return None


# === [新增] 文档库 路由区域 ===

@app.route('/library/view')
def library_view():
    """渲染新的独立文档库页面"""
    if not g.user: return redirect(url_for('login'))
    return render_template('library/index.html', user=g.user)


@app.route('/api/library/files')
def api_library_files():
    """
    文档库专用接口：支持 分类、年份、课程、模糊搜索 组合筛选
    """
    if not g.user: return jsonify({"msg": "Unauthorized"}), 401

    # 获取筛选参数
    category = request.args.get('category', '')  # exam, standard, etc.
    year = request.args.get('year', '')
    course = request.args.get('course', '')
    search = request.args.get('q', '')

    # 调用 Database 的高级筛选方法 (需确保 database.py 中 get_files_by_filter 逻辑正确)
    files = db.get_files_by_filter(
        user_id=g.user['id'],
        doc_category=category if category != 'all' else None,
        year=year,
        course=course,
        # 这里假设 get_files_by_filter 还没支持 search，我们将在 database.py 中微调
    )

    # 如果 database 层没做 search 过滤，这里在内存做 (或者修改 SQL，推荐修改 SQL)
    if search:
        search = search.lower()
        files = [f for f in files if search in f['original_name'].lower()]

    # 标记权限
    for f in files:
        f['is_owner'] = (f['uploaded_by'] == g.user['id']) or g.user.get('is_admin')

    return jsonify(files)


@app.route('/api/library/filters')
def api_library_filters():
    """获取侧边栏筛选树所需的数据 (学年/课程/标签)"""
    if not g.user: return jsonify([]), 401
    tree = db.get_document_library_tree(g.user['id'])
    return jsonify(tree)


@app.route('/export_page/<int:file_id>')
def export_page(file_id):
    if not g.user: return redirect(url_for('login'))

    record = db.get_file_by_id(file_id)
    if not record: return "文件不存在", 404

    # 尝试智能提取课程名 (如果原来的标题里有)
    metadata = {"course_name": "", "class_name": ""}
    if record['original_name']:
        metadata['course_name'] = record['original_name'].replace('.txt', '').replace('.md', '')

    return render_template('export.html', file=record, content=record.get('parsed_content', ''), metadata=metadata,
                           user=g.user)


# ================= 签名集 API =================

@app.route('/api/signatures/list')
def list_signatures():
    if not g.user: return jsonify({"msg": "Unauthorized"}), 401
    search = request.args.get('q', '')
    sigs = db.get_signatures(search)
    # 标记 ownership
    for s in sigs:
        s['is_owner'] = (s['uploaded_by'] == g.user['id']) or (g.user.get('is_admin'))
    return jsonify(sigs)


@app.route('/api/signatures/upload', methods=['POST'])
def upload_signature():
    if not g.user:
        return jsonify({"msg": "Unauthorized"}), 401

    file = request.files.get('file')
    name = request.form.get('name', '').strip()

    if not file: return jsonify({"msg": "请选择文件"}), 400
    if not name: name = file.filename

    # 格式校验
    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in ['.png', '.jpg', '.jpeg']:
        return jsonify({"msg": "仅支持 PNG/JPG 格式"}), 400

    # 物理保存 (Hash去重)
    f_hash = calculate_file_hash(file)
    save_name = f"{f_hash}{ext}"
    save_path = os.path.join(app.config['SIGNATURES_FOLDER'], save_name)

    if not os.path.exists(save_path):
        file.save(save_path)

    # 数据库记录
    db.add_signature(name, f_hash, save_path, g.user['id'])
    return jsonify({"status": "success"})


@app.route('/api/signatures/delete', methods=['POST'])
def delete_signature():
    if not g.user:
        return jsonify({"msg": "Unauthorized"}), 401
    sig_id = request.json.get('id')

    sig = db.get_signature_by_id(sig_id)
    if not sig:
        return jsonify({"msg": "记录不存在"}), 404

    # 权限检查
    if sig['uploaded_by'] != g.user['id'] and not g.user.get('is_admin'):
        return jsonify({"msg": "无权删除他人上传的签名"}), 403

    # 逻辑删除：先删DB记录
    db.delete_signature(sig_id)

    # [修复] 清理物理文件 (使用路径修复逻辑)
    usage = db.get_signature_usage_count(sig['file_hash'])

    # 获取真实路径
    real_path = get_corrected_path(sig['file_path'], app.config['SIGNATURES_FOLDER'])

    if usage == 0 and real_path and os.path.exists(real_path):
        try:
            os.remove(real_path)
        except Exception as e:
            print(f"Failed to remove signature file: {e}")

    return jsonify({"status": "success"})


@app.route('/api/signatures/image/<int:sig_id>')
def get_signature_image(sig_id):
    if not g.user:
        return jsonify({"msg": "Unauthorized"}), 401
    sig = db.get_signature_by_id(sig_id)

    if not sig:
        return "Record not found", 404

    # [修复] 核心修复点：获取真实路径
    real_path = get_corrected_path(sig['file_path'], app.config['SIGNATURES_FOLDER'])

    if not real_path:
        # 如果还是找不到，说明文件可能真的没上传到服务器
        return "Image file not found on server", 404

    return send_file(real_path)


@app.route('/intro')
def intro_page():
    return render_template('intro.html')


@app.route('/api/create_direct_grader', methods=['POST'])
def create_direct_grader():
    if not g.user:
        return jsonify({"msg": "Unauthorized"}), 401

    name = request.form.get('grader_name')
    extra_instruction = request.form.get('extra_instruction', '')

    f_exam = request.files.get('exam_file')
    exam_file_id = request.form.get('exam_file_id')

    f_std = request.files.get('standard_file')
    std_file_id = request.form.get('standard_file_id')

    try:
        # 处理文件保存/查重
        exam_path, _ = handle_file_upload_or_reuse(f_exam, exam_file_id, g.user['id'])
        std_path, _ = handle_file_upload_or_reuse(f_std, std_file_id, g.user['id'])

        if not exam_path or not std_path:
            return jsonify({"msg": "请上传完整的试卷和评分标准"}), 400

        # 获取 ID 用于解析
        exam_record = db.get_file_by_hash(calculate_file_hash(open(exam_path, 'rb')))
        std_record = db.get_file_by_hash(calculate_file_hash(open(std_path, 'rb')))

        # 智能解析 (AI Parsing)
        exam_sta, parsed_exam, _ = smart_parse_file_content(exam_record['id'])
        std_sta, parsed_std, _ = smart_parse_file_content(std_record['id'])

        if not exam_sta:
            return jsonify({"msg": f"试卷解析失败: {parsed_exam}"}), 400
        if not std_sta:
            return jsonify({"msg": f"评分标准解析失败: {parsed_std}"}), 400

        # 生成代码
        grader_id = f"direct_{uuid.uuid4().hex[:8]}"
        class_name = f"DirectGrader_{grader_id}"

        # === 【核心修复】 转义安全函数 ===
        # 即使解析结果中包含 Windows 路径 (如 C:\Users\...)，经过这里处理后也会变为安全的字符串
        def safe_str(s):
            if not s: return ""

            # 1. 【新增】去除 Null Bytes (空字节)，防止二进制乱码导致 "source code string cannot contain null bytes"
            s = s.replace('\x00', '')

            # 2. 先转义反斜杠
            s = s.replace('\\', '\\\\')

            # 3. 再转义引号
            return s.replace('"', '\\"').replace("'''", "\\'\\'\\'")

        code = DIRECT_GRADER_TEMPLATE.format(
            class_name=class_name,
            grader_id=grader_id,
            display_name=name,
            exam_content=safe_str(parsed_exam),
            std_content=safe_str(parsed_std),
            extra_instruction=safe_str(extra_instruction)
        )

        save_path = os.path.join(app.config['GRADERS_DIR'], f"{grader_id}.py")
        with open(save_path, 'w', encoding='utf-8') as f:
            f.write(code)

        db.insert_ai_task(
            name=name, status='success',
            log_info='Direct AI Core Created (Parsed & Cached)',
            exam_path=exam_path, standard_path=std_path,
            strictness='direct', user_id=g.user['id'],
            grader_id=grader_id
        )

        GraderFactory._loaded = False
        GraderFactory.load_graders()

        return jsonify({"status": "success"})

    except Exception as e:
        traceback.print_exc()
        return jsonify({"msg": f"创建失败: {str(e)}"}), 500


@app.route('/api/mark_help_read', methods=['POST'])
def mark_help_read():
    if not g.user:
        return jsonify({"msg": "Unauthorized"}), 401
    db.mark_help_read(g.user['id'])
    return jsonify({"msg": "OK"})


@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form.get('username')
        if not username:
            return render_template('login.html', error="请输入用户名")
        user = db.login_simple_user(username)
        session['user'] = user
        return redirect(url_for('index'))
    return render_template('login.html')


@app.route('/logout')
def logout():
    session.pop('user', None)
    return redirect(url_for('login'))


@app.route('/')
def index():
    if not g.user:
        return redirect(url_for('login'))
    classes = db.get_classes(user_id=g.user['id'])
    return render_template('index.html', classes=classes, user=g.user)


@app.route('/api/my_files')
def my_files():
    if not g.user: return jsonify({"msg": "Unauthorized"}), 401
    search = request.args.get('q', '')
    files = db.get_user_recent_files(g.user['id'], limit=50, search_name=search)
    return jsonify(files)


@app.route('/api/files')
def all_files():
    search = request.args.get('q', '')
    files = db.get_files(limit=50, search_name=search)
    return jsonify(files)


@app.route('/ai_generator')
def ai_generator_page():
    ref_task_id = request.args.get('ref_task_id')
    ref_task = None
    if ref_task_id:
        ref_task = db.get_ai_task_by_id(ref_task_id)

    GraderFactory.load_graders()
    strategies = GraderFactory.get_all_strategies()
    display_list = []

    active_ids = set()
    for sid, sname in strategies:
        active_ids.add(sid)
        task = db.get_task_by_grader_id(sid)
        # 只有 Admin 或者 创建者本人 才是 is_owner
        is_owner = False
        if task and task.get('creator_id') == g.user['id']:
            is_owner = True

        creator_name = task.get('creator_name', 'System') if task else 'Built-in'
        info = {
            "type": "grader",
            "id": sid,
            "name": sname,
            "status": "success",
            "log_info": "Ready",
            "created_at": task['created_at'] if task else "未知",
            "source": "AI 生成" if task else "系统内置",
            "creator": creator_name,
            "is_owner": is_owner
        }
        display_list.append(info)

    recent_tasks = db.get_ai_tasks(limit=50)

    for t in recent_tasks:
        if t['status'] == 'deleted':
            continue

        is_owner = (t.get('created_by') == g.user['id'])

        # === 【修改点 1】: 过滤他人的失败/进行中任务，保持列表干净 ===
        # 如果任务失败，且不是我创建的，则不显示
        if t['status'] == 'failed' and not is_owner:
            continue

        if t['status'] in ['pending', 'processing', 'failed']:
            display_list.append({
                "type": "task",
                "id": t['id'],
                "task_name": t['name'],
                "status": t['status'],
                "log_info": t['log_info'],
                "created_at": t['created_at'],
                "source": "正在生成..." if t['status'] != 'failed' else "生成失败",
                "creator": t.get('creator_name', 'Unknown'),
                "is_owner": is_owner,
                "grader_id": t.get('grader_id')  # 确保传递 grader_id 以便删除
            })
        elif t['status'] == 'success' and t['grader_id'] and t['grader_id'] not in active_ids:
            # 数据库显示成功，但内存没加载到，提示刷新
            expected_path = os.path.join(Config.GRADERS_DIR, f"{t['grader_id']}.py")
            if os.path.exists(expected_path):
                display_list.append({
                    "type": "grader",
                    "id": t['grader_id'],
                    "name": t['name'] + " (需刷新)",
                    "status": "success",
                    "log_info": "已就绪",
                    "created_at": t['created_at'],
                    "source": "AI 生成",
                    "creator": t.get('creator_name', 'Unknown'),
                    "is_owner": is_owner
                })
            else:
                # 文件彻底丢了，标为 deleted
                db.update_task_status_by_grader_id(t['grader_id'], "deleted")
                continue

    display_list.sort(key=lambda x: x['created_at'], reverse=True)

    return render_template('ai_generator.html', ref_task=ref_task, graders=display_list, user=g.user)


@app.route('/grader/<string:grader_id>')
def grader_detail(grader_id):
    GraderFactory.load_graders()
    grader_cls = GraderFactory._graders.get(grader_id)
    task_info = db.get_task_by_grader_id(grader_id)
    file_path = os.path.join(Config.GRADERS_DIR, f"{grader_id}.py")
    code_content = ""
    file_exists = os.path.exists(file_path)

    if file_exists:
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                code_content = f.read()
        except:
            code_content = "# 读取文件失败"

    if not grader_cls and not task_info and not file_exists:
        return "核心未找到或已被删除", 404

    if grader_cls:
        display_name = grader_cls.NAME
        status_tag = "normal"
    else:
        base_name = task_info.get('name', grader_id) if task_info else grader_id
        display_name = f"{base_name} (⚠ 加载失败-请检查代码)"
        status_tag = "error"

    is_owner = False
    if task_info and g.user and task_info.get('creator_id') == g.user['id']:
        is_owner = True
    if not task_info and g.user and g.user.get('is_admin'):
        is_owner = True

    return render_template('grader_detail.html',
                           grader={"id": grader_id, "name": display_name, "code": code_content, "status": status_tag},
                           task=task_info, is_owner=is_owner, user=g.user)


@app.route('/api/create_grader_task', methods=['POST'])
def create_grader_task():
    name = request.form.get('task_name', '未命名任务')
    ref_task_id = request.form.get('ref_task_id')
    strictness = request.form.get('strictness', 'standard')
    extra_desc = request.form.get('extra_desc', '')
    try:
        max_score = int(request.form.get('max_score', 100))
    except:
        max_score = 100

    f_exam = request.files.get('exam_file')
    f_std = request.files.get('standard_file')
    exam_file_id = request.form.get('exam_file_id')
    std_file_id = request.form.get('standard_file_id')

    try:
        exam_path, exam_name = handle_file_upload_or_reuse(f_exam, exam_file_id, g.user['id'])
        std_path, std_name = handle_file_upload_or_reuse(f_std, std_file_id, g.user['id'])

        if ref_task_id:
            old_task = db.get_ai_task_by_id(ref_task_id)
            if old_task:
                if not exam_path and old_task.get('exam_path') and os.path.exists(old_task['exam_path']):
                    exam_path = old_task['exam_path']
                if not std_path and old_task.get('standard_path') and os.path.exists(old_task['standard_path']):
                    std_path = old_task['standard_path']

        if not exam_path or not std_path:
            return jsonify({"msg": "请提供完整的试卷和评分标准"}), 400

        # === 新增：解析校验逻辑 ===
        exa_sta, exam_text = extract_text_from_file(exam_path)
        std_sta, std_text = extract_text_from_file(std_path)

        # 检查是否解析失败
        if not exa_sta or exam_text.startswith("解析失败") or exam_text.startswith("解析结果为空"):
            return jsonify({"msg": f"试卷解析失败: {exam_text}。请尝试手动粘贴内容。"}), 400

        if not std_sta or std_text.startswith("解析失败") or std_text.startswith("解析结果为空"):
            return jsonify({"msg": f"评分标准解析失败: {std_text}。请尝试手动粘贴内容。"}), 400

        task_id = db.insert_ai_task(name, "pending", "任务已提交...",
                                    exam_path, std_path,
                                    strictness, extra_desc, max_score,
                                    user_id=g.user['id'])

        t = threading.Thread(target=ai_generation_worker,
                             args=(task_id, exam_text, std_text, strictness, extra_desc, max_score))
        t.start()

        return jsonify({"msg": "任务已提交", "task_id": task_id})

    except Exception as e:
        traceback.print_exc()
        return jsonify({"msg": f"处理失败: {str(e)}"}), 500


@app.route('/api/get_ai_tasks')
def get_ai_tasks():
    tasks = db.get_ai_tasks(limit=10)
    return jsonify([dict(t) for t in tasks])


@app.route('/api/save_grader_code', methods=['POST'])
def save_grader_code():
    data = request.json
    grader_id = data.get('id')
    new_code = data.get('code')
    task = db.get_task_by_grader_id(grader_id)
    if not task or task.get('creator_id') != g.user['id']:
        return jsonify({"msg": "您无权修改他人创建的核心代码"}), 403

    if not grader_id or not new_code: return jsonify({"msg": "参数错误"}), 400

    file_path = os.path.join(Config.GRADERS_DIR, f"{grader_id}.py")
    if not os.path.exists(file_path): return jsonify({"msg": "文件不存在"}), 404

    try:
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(new_code)
        GraderFactory._loaded = False
        GraderFactory.load_graders()
        return jsonify({"msg": "保存成功并已重载"})
    except Exception as e:
        return jsonify({"msg": f"保存失败: {str(e)}"}), 500


@app.route('/api/delete_grader', methods=['POST'])
def delete_grader():
    # 前端可能传过来 grader_id (str) 或者 task_id (int/str)
    raw_id = request.json.get('id')
    if not raw_id:
        return jsonify({"msg": "ID不能为空"}), 400

    task = None

    # 1. 尝试逻辑：先判断是不是 Task ID (数字)
    # 失败的任务通常只剩下 Task ID，没有 grader_id
    if str(raw_id).isdigit():
        task = db.get_ai_task_by_id(int(raw_id))

    # 2. 如果没找到，或者传来的明显是 grader_id (如 "direct_xxxx")
    # 则尝试按 grader_id 查找
    if not task:
        task = db.get_task_by_grader_id(raw_id)

    if not task:
        # 如果既找不到任务，文件也不存在，可能是旧数据遗留
        return jsonify({"msg": "任务记录不存在，无法删除"}), 404

    # 3. 权限检查
    # 必须是创建者或者管理员
    if task.get('created_by') != g.user['id'] and not g.user.get('is_admin'):
        return jsonify({"msg": "您无权删除他人创建的核心"}), 403

    # 4. 执行删除操作

    # A. 物理文件删除 (只有当 grader_id 存在时才需要删文件)
    grader_id = task.get('grader_id')
    if grader_id:
        file_path = os.path.join(Config.GRADERS_DIR, f"{grader_id}.py")
        # 即使文件不存在也不报错，继续删数据库记录
        if os.path.exists(file_path):
            timestamp = int(time.time())
            backup_name = f"{grader_id}_{timestamp}.py.bak"
            backup_path = os.path.join(Config.TRASH_DIR, backup_name)
            try:
                shutil.move(file_path, backup_path)

                # 记录进回收站表
                GraderFactory.load_graders()
                g_cls = GraderFactory._graders.get(grader_id)
                name = g_cls.NAME if g_cls else (task.get('name') or grader_id)
                db.recycle_grader_record(grader_id, name, backup_name)
            except Exception as e:
                print(f"[Delete Error] Move file failed: {e}")

    # B. 数据库状态软删除 (关键修复：使用 Task ID 更新)
    # 这样即使 grader_id 为空，也能成功将状态置为 deleted
    db.update_ai_task_status(task['id'], "deleted")

    # C. 尝试热重载 (移除内存中的对象)
    GraderFactory._loaded = False
    GraderFactory.load_graders()

    return jsonify({"msg": "已移入回收站"})


@app.route('/api/get_recycled_graders')
def get_recycled():
    items = db.get_recycled_graders()
    return jsonify(items)


@app.route('/api/restore_grader', methods=['POST'])
def restore_grader():
    recycle_id = request.json.get('id')
    record = db.restore_grader_record(recycle_id)
    if record:
        backup_path = os.path.join(Config.TRASH_DIR, record['backup_filename'])
        target_path = os.path.join(Config.GRADERS_DIR, f"{record['grader_id']}.py")
        if os.path.exists(backup_path):
            shutil.move(backup_path, target_path)
            db.update_task_status_by_grader_id(record['grader_id'], "success")
            GraderFactory._loaded = False
            GraderFactory.load_graders()
            return jsonify({"msg": "恢复成功"})
        return jsonify({"msg": "备份文件丢失"}), 404
    return jsonify({"msg": "记录不存在"}), 404


@app.route('/api/prepare_regeneration', methods=['POST'])
def prepare_regeneration():
    grader_id = request.json.get('grader_id')
    task = db.get_task_by_grader_id(grader_id)
    if not task:
        return jsonify({"msg": "无法找到该核心的原始生成任务"}), 404

    # 逻辑同删除
    file_path = os.path.join(Config.GRADERS_DIR, f"{grader_id}.py")
    if os.path.exists(file_path):
        timestamp = int(time.time())
        backup_name = f"{grader_id}_{timestamp}.py.bak"
        shutil.move(file_path, os.path.join(Config.TRASH_DIR, backup_name))

        GraderFactory.load_graders()
        g_cls = GraderFactory._graders.get(grader_id)
        name = g_cls.NAME if g_cls else grader_id
        db.recycle_grader_record(grader_id, name, backup_name)

        GraderFactory._loaded = False
        GraderFactory.load_graders()

    return jsonify({"msg": "OK", "ref_task_id": task['id']})


@app.route('/new_class', methods=['GET', 'POST'])
def new_class():
    if request.method == 'POST':
        cname = request.form['class_name']
        course = request.form['course_name']
        strategy = request.form.get('strategy', '')
        file = request.files['student_list']

        if file:
            cid = db.create_class(cname, course, strategy, g.user['id'])
            class_ws = os.path.join(app.config['WORKSPACE_FOLDER'], str(cid))
            os.makedirs(os.path.join(class_ws, 'raw_zips'), exist_ok=True)
            os.makedirs(os.path.join(class_ws, 'extracted'), exist_ok=True)
            db.update_class_workspace(cid, class_ws)
            try:
                df = pd.read_excel(file) if file.filename.endswith('.xlsx') else pd.read_csv(file)
                df.columns = [c.strip() for c in df.columns]
                sid_col = next((c for c in df.columns if '学号' in c), None)
                name_col = next((c for c in df.columns if '姓名' in c), None)
                if not sid_col or not name_col: return "名单缺少【学号】或【姓名】列", 400
                for _, row in df.iterrows():
                    db.add_student(str(row[sid_col]), str(row[name_col]), cid)
            except Exception as e:
                return f"名单解析失败: {str(e)}", 400
            return redirect(url_for('grading_view', class_id=cid))

    GraderFactory.load_graders()
    strategies = GraderFactory.get_all_strategies()
    default_strategy = request.args.get('strategy', '')
    return render_template('newClass.html', strategies=strategies, default_strategy=default_strategy, user=g.user)


@app.route('/grading/<int:class_id>')
def grading_view(class_id):
    cls = db.get_class_by_id(class_id)
    if not cls: return "Class not found", 404
    if cls['created_by'] != g.user['id']:
        return "无权访问此班级", 403
    students = db.get_students_with_grades(class_id)
    return render_template('grading.html', cls=cls, students=students, user=g.user)


@app.route('/grading/<int:class_id>/student/<string:student_id>')
def student_detail(class_id, student_id):
    cls = db.get_class_by_id(class_id)
    student = db.get_student_detail(class_id, student_id)
    if not student: return "Student not found", 404

    ws = get_real_workspace_path(class_id)
    extract_path = os.path.join(ws, 'extracted', str(student_id))

    def get_file_tree(root_path):
        tree = []
        try:
            items = sorted(os.listdir(root_path))
            for item in items:
                full_path = os.path.join(root_path, item)
                node = {'name': item, 'type': 'folder' if os.path.isdir(full_path) else 'file'}
                if node['type'] == 'folder':
                    node['children'] = get_file_tree(full_path)
                else:
                    node['size'] = os.path.getsize(full_path)
                tree.append(node)
        except Exception:
            pass
        return tree

    file_tree = []
    if os.path.exists(extract_path):
        file_tree = get_file_tree(extract_path)

    zip_info = {"name": "未提交", "size": 0}
    if student['filename']:
        zip_path = os.path.join(ws, 'raw_zips', student['filename'])
        if os.path.exists(zip_path):
            zip_info['name'] = student['filename']
            zip_info['size'] = round(os.path.getsize(zip_path) / 1024, 2)

    return render_template('student_detail.html', cls=cls, s=student, zip_info=zip_info, file_tree=file_tree)


@app.route('/upload_zips/<int:class_id>', methods=['POST'])
def upload_zips(class_id):
    # cls = db.get_class_by_id(class_id) # 这一行其实不需要了，因为我们只用ID

    # 【修复】不再使用 cls['workspace_path']，而是动态生成路径
    workspace_path = get_real_workspace_path(class_id)

    raw_dir = os.path.join(workspace_path, 'raw_zips')
    if not os.path.exists(raw_dir):
        os.makedirs(raw_dir)

    files = request.files.getlist('files')
    count = 0
    for file in files:
        if file and file.filename:
            # 使用 os.path.join 确保路径分隔符正确
            save_path = os.path.join(raw_dir, file.filename)
            file.save(save_path)
            count += 1

    # 【可选】顺便更新一下数据库里的路径，方便以后查看，但这步不是必须的
    try:
        db.update_class_workspace(class_id, workspace_path)
    except:
        pass

    return jsonify({"msg": f"上传 {count} 个文件成功"})


@app.route('/run_grading_logic/<int:class_id>', methods=['POST'])
def run_grading_logic(class_id):
    cls = db.get_class_by_id(class_id)
    conn = db.get_connection()
    students = conn.execute("SELECT * FROM students WHERE class_id=?", (class_id,)).fetchall()

    # 【修复】动态获取真实路径
    workspace_path = get_real_workspace_path(class_id)

    raw_dir = os.path.join(workspace_path, 'raw_zips')
    extract_base = os.path.join(workspace_path, 'extracted')

    # 确保目录存在
    if not os.path.exists(extract_base): os.makedirs(extract_base)

    grader = GraderFactory.get_grader(cls['strategy'])
    if not grader:
        return jsonify({"msg": f"未找到策略 [{cls['strategy']}]"}), 500

    if not os.path.exists(raw_dir): return jsonify({"msg": "请先上传文件"}), 400
    uploaded_files = os.listdir(raw_dir)

    db.clear_grades(class_id)

    for s in students:
        sid = str(s['student_id'])
        name = s['name']
        matched_file = None
        for f in uploaded_files:
            if sid in f or name in f:
                matched_file = f
                break

        if not matched_file:
            db.save_grade_error(sid, class_id, "未找到提交文件", "")
            continue

        student_extract_dir = os.path.join(extract_base, sid)
        if os.path.exists(student_extract_dir): shutil.rmtree(student_extract_dir)
        os.makedirs(student_extract_dir)

        archive_path = os.path.join(raw_dir, matched_file)
        try:
            patoolib.extract_archive(archive_path, outdir=student_extract_dir, verbosity=-1)
            result = grader.grade(student_extract_dir, {"sid": sid, "name": name})
            status = "PASS" if result.is_pass else "FAIL"
            db.save_grade(sid, class_id, result.total_score, result.get_details_json(),
                          result.get_deduct_str(), status, matched_file)
        except Exception as e:
            err_msg = f"解压/批改异常: {str(e)}"
            if "rar" in str(matched_file).lower() and "patool" in str(e):
                err_msg = "解压失败(服务器可能缺少WinRAR)"
            db.save_grade_error(sid, class_id, err_msg, matched_file)

    return jsonify({"msg": "批改完成"})


@app.route('/clear_data/<int:class_id>', methods=['POST'])
def clear_data(class_id):
    cls = db.get_class_by_id(class_id)
    db.clear_grades(class_id)
    workspace = get_real_workspace_path(class_id)
    if os.path.exists(workspace):
        shutil.rmtree(workspace)
        os.makedirs(os.path.join(workspace, 'raw_zips'), exist_ok=True)
        os.makedirs(os.path.join(workspace, 'extracted'), exist_ok=True)
    return jsonify({"msg": "数据已清空"})


@app.route('/delete_class/<int:class_id>', methods=['POST'])
def delete_class(class_id):
    cls = db.get_class_by_id(class_id)
    workspace = get_real_workspace_path(class_id)
    if os.path.exists(workspace):
        shutil.rmtree(workspace)
    db.delete_class(class_id)
    return jsonify({"msg": "批改任务已删除"})


@app.route('/export/<int:class_id>')
def export_excel(class_id):
    import json
    with db.get_connection() as conn:
        students_data = conn.execute('''
                                     SELECT s.student_id,
                                            s.name,
                                            g.total_score,
                                            g.score_details,
                                            g.deduct_details,
                                            g.filename
                                     FROM students s
                                              LEFT JOIN grades g ON s.student_id = g.student_id AND g.class_id = s.class_id
                                     WHERE s.class_id = ?
                                     ''', (class_id,)).fetchall()
        cls = db.get_class_by_id(class_id)

    data_list = []
    for row in students_data:
        item = {
            "学号": row['student_id'],
            "姓名": row['name'],
            "总分": row['total_score'] if row['total_score'] is not None else 0,
            "扣分详情": row['deduct_details'],
            "文件名": row['filename']
        }
        if row['score_details']:
            try:
                details = json.loads(row['score_details'])
                for d in details:
                    col_name = d.get('name', '未知项')
                    item[col_name] = d.get('score', 0)
            except:
                pass
        data_list.append(item)

    df = pd.DataFrame(data_list)
    fixed_cols = ['学号', '姓名']
    end_cols = ['总分', '扣分详情', '文件名']
    dynamic_cols = [c for c in df.columns if c not in fixed_cols and c not in end_cols]
    final_cols = fixed_cols + dynamic_cols + end_cols
    df = df[final_cols]

    path = os.path.join(app.config['UPLOAD_FOLDER'], f"{cls['name']}_成绩单.xlsx")
    df.to_excel(path, index=False)
    return send_file(path, as_attachment=True)


@app.route('/file_manager')
def file_manager_page():
    if not g.user: return redirect(url_for('login'))
    return render_template('file_manager.html', user=g.user)


@app.route('/api/parse_file_asset', methods=['POST'])
def parse_file_asset():
    if not g.user:
        return jsonify({"msg": "Unauthorized"}), 401

    file_obj = request.files.get('file')
    if not file_obj or not file_obj.filename:
        return jsonify({"msg": "请先选择文件"}), 400

    try:
        physical_path, _ = handle_file_upload_or_reuse(file_obj, None, g.user['id'])
        if not physical_path:
            return jsonify({"msg": "文件保存失败"}), 400

        with open(physical_path, 'rb') as f:
            record = db.get_file_by_hash(calculate_file_hash(f))

        if not record:
            return jsonify({"msg": "文件记录创建失败"}), 500

        # parse_success, parsed_text, _ = smart_parse_file_content(record['id'], request.form.get('doc_type'))
        doc_type = request.form.get('doc_type')
        parse_success, parsed_text, _ = smart_parse_with_metadata(record['id'], doc_type)
        if not parse_success:
            return jsonify({"msg": parsed_text}), 400

        return jsonify({
            "status": "success",
            "file_id": record['id'],
            "title": record['original_name'],
            "parsed_content": parsed_text
        })
    except Exception as e:
        traceback.print_exc()
        return jsonify({"msg": f"解析失败: {str(e)}"}), 500


@app.route('/api/save_pasted_document', methods=['POST'])
def save_pasted_document():
    if not g.user:
        return jsonify({"msg": "Unauthorized"}), 401

    data = request.json or {}
    content = (data.get('content') or '').strip()
    doc_type = data.get('doc_type')

    if not content:
        return jsonify({"msg": "内容不能为空"}), 400

    title = generate_title_with_ai(content, doc_type)
    file_id, filename = create_text_asset(content, title, g.user['id'])
    if not file_id:
        return jsonify({"msg": "保存失败"}), 500

    return jsonify({"status": "success", "file_id": file_id, "title": filename})


@app.route('/api/parse_and_save_pasted_document', methods=['POST'])
def parse_and_save_pasted_document():
    """
    粘贴入库：先让 AI 整理成 JSON 结构，然后后端拆包存库。
    解决：之前存入的是 raw json 导致的格式错误问题。
    """
    if not g.user:
        return jsonify({"msg": "Unauthorized"}), 401

    data = request.json or {}
    content = (data.get('content') or '').strip()
    doc_type = data.get('doc_type')

    if not content:
        return jsonify({"msg": "内容不能为空"}), 400

    doc_label = "试卷" if doc_type == "exam" else "评分标准"
    standard_config = db.get_best_ai_config("standard") or db.get_best_ai_config("thinking")

    if not standard_config:
        # 如果没有 AI 配置，直接保存原始内容
        title = generate_title_from_content(content, doc_type)
        file_id, filename = create_text_asset(content, title, g.user['id'])
        return jsonify({"status": "success", "file_id": file_id, "title": filename, "msg": "未配置AI，已原样保存"})

    # Prompt 要求返回 JSON，以便分离标题和正文
    prompt = (
        f"请对以下{doc_label}内容进行深度规整。\n"
        "1. 去除无用的页眉页脚、乱码。\n"
        "2. 将内容整理为清晰的 Standard Markdown 格式。\n"
        "3. 根据文档类型和内容为文档起一个合理的标题。\n"
        "4. **必须以纯 JSON 格式返回**，不要使用代码块标记，格式如下：\n"
        "{\"title\": \"文档标题\", \"content\": \"Markdown格式的正文内容...\"}\n\n"
        f"原始内容：\n{content}"
    )

    try:
        ai_text = asyncio.run(call_ai_platform_chat(
            system_prompt="你是文档整理助手。你只输出 JSON。",
            messages=[{"role": "user", "content": prompt}],
            platform_config=standard_config
        ))

        # 使用增强版工具函数提取
        title, normalized_content = extract_title_and_content(ai_text, doc_type)

        if not normalized_content:
            return jsonify({"msg": "AI 解析内容为空"}), 500

        # === 核心修正：这里存入的是 normalized_content (Markdown)，而不是 ai_text (JSON) ===
        file_id, filename = create_text_asset(normalized_content, title, g.user['id'])

        if not file_id:
            return jsonify({"msg": "保存失败"}), 500

        return jsonify({"status": "success", "file_id": file_id, "title": filename})
    except Exception as e:
        traceback.print_exc()
        return jsonify({"msg": f"解析保存失败: {str(e)}"}), 500


@app.route('/api/update_file_content', methods=['POST'])
def update_file_content():
    if not g.user:
        return jsonify({"msg": "Unauthorized"}), 401

    data = request.json or {}
    file_id = data.get('id')
    content = data.get('content')

    if not file_id or content is None:
        return jsonify({"msg": "参数错误"}), 400

    record = db.get_file_by_id(file_id)
    if not record:
        return jsonify({"msg": "文件不存在"}), 404

    # 权限检查：只能编辑自己的文件，或者是管理员
    # 注意：record['uploaded_by'] 可能是 int，g.user['id'] 也是 int，建议强转比较
    if int(record['uploaded_by']) != int(g.user['id']) and not g.user.get('is_admin'):
        return jsonify({"msg": "无权编辑他人文档"}), 403

    # 调用数据库更新方法
    db.update_file_parsed_content(file_id, content)
    return jsonify({"status": "success", "msg": "保存成功"})


@app.route('/api/ai_generate_document', methods=['POST'])
def ai_generate_document():
    if not g.user:
        return jsonify({"msg": "Unauthorized"}), 401

    prompt = (request.form.get('prompt') or '').strip()
    doc_type = request.form.get('doc_type') or 'exam'

    # 获取上传的新文件
    uploaded_files = request.files.getlist('files')
    # 获取用户选择的已有文件 ID
    existing_file_ids = request.form.getlist('existing_file_ids')

    if not prompt and not uploaded_files and not existing_file_ids:
        return jsonify({"msg": "请输入提示词或提供参考素材"}), 400

    image_payloads = []
    doc_texts = []
    file_errors = []

    try:
        # A. 处理新上传的文件 (保持不变)
        for f in uploaded_files:
            if not f or not f.filename: continue
            physical_path, original_name = handle_file_upload_or_reuse(f, None, g.user['id'])
            if not physical_path: continue

            ext = os.path.splitext(physical_path)[1].lower()
            if ext in ['.png', '.jpg', '.jpeg', '.bmp', '.webp', '.tiff']:
                with open(physical_path, 'rb') as img_f:
                    b64 = base64.b64encode(img_f.read()).decode('utf-8')
                mime_type = mimetypes.types_map.get(ext, 'image/png')
                image_payloads.append({"type": "image", "data": f"data:{mime_type};base64,{b64}"})
            else:
                ok, text = extract_text_from_file(physical_path)
                if ok:
                    doc_texts.append({"name": original_name or f.filename, "content": text})
                else:
                    file_errors.append(f"{original_name or f.filename}: {text}")

        # B. 处理已选择的现有文件 (保持不变)
        for eid in existing_file_ids:
            if not eid: continue
            record = db.get_file_by_id(eid)
            if record and record.get('parsed_content'):
                doc_texts.append({
                    "name": f"[库]{record['original_name']}",
                    "content": record['parsed_content']
                })

        # C. 组装 Prompt (【修改点】：使用 DocumentTypeConfig)
        # 获取特定类型的 System Prompt (包含 JSON 结构要求)
        system_instruction = DocumentTypeConfig.get_prompt_by_type(doc_type)

        user_message_parts = []
        user_message_parts.append(
            f"【生成任务】\n请根据以下要求和参考素材，生成一份新的【{DocumentTypeConfig.TYPES.get(doc_type, '文档')}】。")
        user_message_parts.append(f"用户具体要求：{prompt}")

        if doc_texts:
            doc_section = ["\n【参考素材内容】："]
            for doc in doc_texts:
                content_snippet = doc['content'][:30000]  # 防止 token 溢出
                doc_section.append(f"--- {doc['name']} ---\n{content_snippet}\n")
            user_message_parts.append("\n".join(doc_section))

        if file_errors:
            user_message_parts.append("\n部分文件解析失败忽略：" + "; ".join(file_errors))

        user_message_parts.append("\n请务必严格遵循 System Prompt 中的 JSON 输出格式，生成完整的 metadata 和 content。")

        final_user_prompt = "\n".join(user_message_parts)

        # D. 调用 AI
        payload = {
            "system_prompt": system_instruction,  # 使用统一的配置
            "messages": [],
            "new_message": final_user_prompt,
            "model_capability": "vision" if image_payloads else "standard"  # 生成任务建议用 standard 或 thinking，除非有图片
        }

        # 针对复杂生成任务，如果配置了 Thinking 模型效果更好
        thinking_config = db.get_best_ai_config("thinking")
        if thinking_config:
            payload["model_capability"] = "thinking"

        if image_payloads:
            payload["messages"].append({
                "role": "user",
                "content": "参考图片如下",
                "file_ids": image_payloads
            })
            payload["model_capability"] = "vision"  # 有图必须用 vision

        response = httpx.post(app.config['AI_ASSISTANT_CHAT_ENDPOINT'], json=payload, timeout=240.0)
        if response.status_code != 200:
            return jsonify({"msg": f"AI 助手返回错误: {response.status_code} - {response.text}"}), 500

        ai_text = response.json().get("response_text", "")

        # E. 解析返回的 JSON (复用 smart_parse 的逻辑)
        # 因为现在 Prompt 强制返回 JSON {metadata:..., content:...}
        try:
            cleaned_json = ai_text.strip()
            # 尝试提取 JSON 块
            import re
            json_match = re.search(r'```(?:json)?\s*(\{.*?\})\s*```', cleaned_json, re.DOTALL)
            if json_match:
                cleaned_json = json_match.group(1)
            else:
                # 尝试找首尾大括号
                start = cleaned_json.find('{')
                end = cleaned_json.rfind('}')
                if start != -1 and end != -1:
                    cleaned_json = cleaned_json[start:end + 1]

            data = json.loads(cleaned_json)

            parsed_content = data.get("content", "")
            meta = data.get("metadata", {})

            # 如果解析失败或者 AI 没遵循格式，回退到纯文本
            if not parsed_content and not meta:
                parsed_content = ai_text  # 整个当做正文
                title = generate_title_from_content(ai_text, doc_type)
            else:
                # 尝试从 Meta 中获取更准确的标题
                title = meta.get("course_name", "") + " " + DocumentTypeConfig.TYPES.get(doc_type, "文档")
                if len(title) < 5: title = generate_title_from_content(parsed_content, doc_type)

        except Exception as e:
            print(f"JSON Parse Error in Generate: {e}")
            parsed_content = ai_text
            title = generate_title_from_content(ai_text, doc_type)
            meta = {}

        # 保存结果
        file_id, filename = create_text_asset(parsed_content, title, g.user['id'])

        # 【关键】更新元数据到数据库
        if file_id and meta:
            conn = db.get_connection()
            conn.execute("UPDATE file_assets SET meta_info = ? WHERE id = ?",
                         (json.dumps(meta, ensure_ascii=False), file_id))
            conn.commit()

        if not file_id:
            return jsonify({"msg": "保存 AI 生成内容失败"}), 500

        return jsonify({"status": "success", "file_id": file_id, "title": filename})

    except Exception as e:
        traceback.print_exc()
        return jsonify({"msg": f"生成失败: {str(e)}"}), 500


@app.route('/api/my_parsed_files')
def my_parsed_files():
    if not g.user: return jsonify({"msg": "Unauthorized"}), 401
    files = db.get_user_parsed_files(g.user['id'])
    for f in files:
        f['is_owner'] = f.get('uploaded_by') == g.user['id']
    return jsonify(files)


@app.route('/api/delete_file_asset', methods=['POST'])
def delete_file_asset():
    if not g.user: return jsonify({"msg": "Unauthorized"}), 401
    file_id = request.json.get('id')
    file_record = db.get_file_by_id(file_id)
    if not file_record: return jsonify({"msg": "文件不存在"}), 404
    if file_record['uploaded_by'] != g.user['id'] and not g.user.get('is_admin'):
        return jsonify({"msg": "无权删除此文件"}), 403
    db.delete_file_asset(file_id)
    try:
        if os.path.exists(file_record['physical_path']):
            os.remove(file_record['physical_path'])
    except:
        pass
    return jsonify({"msg": "删除成功"})


# 2. 新增路由：单人批改接口
@app.route('/api/grade_student/<int:class_id>/<string:student_id>', methods=['POST'])
def api_grade_single_student(class_id, student_id):
    """前端逐个调用的接口"""
    try:
        success, msg, data = _grade_single_student_internal(class_id, student_id)
        if success:
            return jsonify({"status": "success", "msg": msg, "data": data})
        else:
            return jsonify({"status": "error", "msg": msg, "filename": data if isinstance(data, str) else ""})
    except Exception as e:
        return jsonify({"status": "error", "msg": str(e)}), 500


# === 添加文件预览路由 ===
@app.route('/preview_file/<int:class_id>/<string:student_id>')
def preview_file(class_id, student_id):
    if not g.user:
        return jsonify({"msg": "Unauthorized"}), 401

    # 1. 获取请求的相对路径
    # 前端传参: ?path=文件夹/文件名.py
    rel_path = request.args.get('path', '')
    if not rel_path:
        return jsonify({"msg": "Path required"}), 400

    # 2. 构建绝对路径
    # 使用与 student_detail 相同的路径逻辑
    ws_path = get_real_workspace_path(class_id)
    base_dir = os.path.join(ws_path, 'extracted', str(student_id))

    # 安全拼接路径 (防止 ../ 目录遍历攻击)
    # normpath 会把 a/b/../c 变成 a/c
    full_path = os.path.normpath(os.path.join(base_dir, rel_path))

    # 安全检查：确保拼接后的路径依然在 base_dir 目录下
    if not full_path.startswith(os.path.abspath(base_dir)):
        return jsonify({"msg": "Illegal path access"}), 403

    if not os.path.exists(full_path):
        return jsonify({"msg": "File not found"}), 404

    if os.path.isdir(full_path):
        return jsonify({"msg": "Cannot preview a directory"}), 400

    # 3. 判断文件类型并返回
    # A. 如果是图片，直接返回文件流，前端会以 Blob 形式加载
    mime_type, _ = mimetypes.guess_type(full_path)
    if mime_type and mime_type.startswith('image'):
        return send_file(full_path, mimetype=mime_type)

    # B. 如果是文本/代码，读取内容并以 JSON 返回
    # 这里复用你已有的 extract_text_from_file 或简单的文本读取逻辑
    try:
        # 尝试以文本方式读取 (尝试 utf-8, gbk 等)
        content = ""
        encodings = ['utf-8', 'gbk', 'gb18030', 'latin1']
        for enc in encodings:
            try:
                with open(full_path, 'r', encoding=enc) as f:
                    content = f.read()
                break
            except UnicodeDecodeError:
                continue

        # 如果读取成功，返回 JSON
        if content:
            # 简单的二进制检测，防止返回乱码
            if '\0' in content:
                return jsonify({"type": "binary", "msg": "Binary file detected", "size": os.path.getsize(full_path)})

            return jsonify({
                "type": "text",
                "content": content,
                "size": os.path.getsize(full_path)
            })
        else:
            # 如果尝试了所有编码都读不出，大概率是二进制文件
            return jsonify({"type": "binary", "msg": "Unable to decode text", "size": os.path.getsize(full_path)})

    except Exception as e:
        print(f"Preview Error: {e}")
        return jsonify({"type": "error", "msg": str(e)}), 500


# 2. 获取可用模板列表 API
@app.route('/api/export/templates')
def get_export_templates():
    if not g.user: return jsonify({"msg": "Unauthorized"}), 401
    # 从数据库读取，确保是最新的
    conn = db.get_connection()
    templates = conn.execute(
        "SELECT template_id, name, description, ui_schema FROM export_templates WHERE is_active=1").fetchall()
    return jsonify([dict(row) for row in templates])


# 3. 执行导出 API (重写)
# 需要在 app.py 顶部确保引入 quote
from urllib.parse import quote


@app.route('/api/export_word_v2', methods=['POST'])
def export_word_v2():
    if not g.user: return jsonify({"msg": "Unauthorized"}), 401

    data = request.json
    file_id = data.get('file_id')
    template_id = data.get('template_id')
    form_data = data.get('form_data', {})  # 前端动态表单填写的KV

    # 获取文件内容和元数据
    file_record = db.get_file_by_id(file_id)
    if not file_record:
        return jsonify({"msg": "文件记录不存在"}), 404

    content = file_record['parsed_content']
    meta_info = json.loads(file_record.get('meta_info') or '{}')

    # 获取模板实例
    exporter = TemplateManager.get_template(template_id)
    if not exporter:
        return jsonify({"msg": "模板未找到"}), 404

    # === [FIX 1] 修复字典遍历报错 ===
    # 处理签名路径逻辑 (将 ID 转为 Path)
    # 使用 list() 创建 items 的副本进行遍历，避免 "dictionary changed size" 错误
    for key, val in list(form_data.items()):
        # 只要 key 是 _sig 结尾（如 teacher_sig）或者是特定的 _select 结尾（兼容旧逻辑），且值为数字 ID
        is_sig_field = key.endswith('_sig') or key.endswith('_select')

        if is_sig_field and str(val).isdigit():
            sig = db.get_signature_by_id(int(val))
            if sig:
                # 注入 path 供模板使用
                # 注意：这里我们向 form_data 添加了新 Key，所以外层必须用 list() 拷贝迭代
                real_path = get_corrected_path(sig['file_path'], app.config['SIGNATURES_FOLDER'])
                if real_path:
                    # 如果 key 是 teacher_sig，则注入 teacher_sig_path
                    # 如果 key 是 teacher_sig_select，则注入 teacher_sig_path (保持兼容)
                    prefix = key.replace('_select', '')
                    form_data[f"{prefix}_path"] = real_path

    # 生成文件
    filename = f"export_{uuid.uuid4().hex}.docx"
    save_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)

    try:
        exporter.generate(content, meta_info, form_data, save_path)

        # === [FIX 2] 补全中文文件名处理 ===
        dl_name = form_data.get('course_name', 'Document')
        # 确保文件名是字符串
        if not dl_name: dl_name = 'Document'

        try:
            # 尝试编码，如果不报错说明是纯 ASCII，不需要 quote
            dl_name.encode('latin-1')
            final_name = f"{dl_name}.docx"
        except UnicodeEncodeError:
            # 包含中文，使用 URL 编码防止 header 错误
            final_name = f"{quote(dl_name)}.docx"

        return send_file(save_path, as_attachment=True, download_name=final_name)

    except Exception as e:
        traceback.print_exc()
        return jsonify({"msg": f"导出失败: {str(e)}"}), 500


if __name__ == '__main__':
    app.run(debug=True, port=5010)
